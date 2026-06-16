"""Extract text from BLIIOT product docx files."""
import docx
import sys
import os

path = sys.argv[1]
doc = docx.Document(path)
lines = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        lines.append(t)

# Also extract tables
for i, table in enumerate(doc.tables):
    lines.append(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        lines.append(" | ".join(cells))

text = "\n".join(lines)
print(text)