"""Batch extract all BLIIOT product datasheets."""
import docx
import os
import json
import sys

base = r'C:\Users\Admin\Desktop\Working\产品规格书\英文资料'

# All key docx files to extract (datasheets and user manuals)
key_files = {
    # ===== ARM嵌入式控制器 =====
    'BL301': ['ARM嵌入式控制器/BL301/说明书/BLIIoT BL301 BL302 User Manual V1.2.docx'],
    'BL302': ['ARM嵌入式控制器/BL302/说明书/BLIIoT BL301 BL302 User Manual V1.2.docx'],
    'BL304': ['ARM嵌入式控制器/BL304/说明书/BLIIoT BL303 BL304 User Manual V1.0.docx'],
    'BL310': [
        'ARM嵌入式控制器/BL310/ARMxy BL310 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL310/BLIIOT BL310 User Manual V1.0.docx',
    ],
    'BL330': [
        'ARM嵌入式控制器/BL330/ARMxy BL330 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL330/BLIIoT BL330 User Manual V1.0.docx',
    ],
    'BL335': [
        'ARM嵌入式控制器/BL335/ARMxy BL335 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL335/BLIIOT ARMxy BL335 User Manual V1.0.docx',
    ],
    'BL340': [
        'ARM嵌入式控制器/BL340/ARMxy BL340 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL340/BLIIoT BL340 User Manual V1.1.docx',
    ],
    'BL350': [
        'ARM嵌入式控制器/BL350/ARMxy BL350 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL350/BLIIOT BL350 User Manual V1.0.docx',
    ],
    'BL360': ['ARM嵌入式控制器/BL360/ARMxy BL360 Datasheet V1.0.docx'],
    'BL370': [
        'ARM嵌入式控制器/BL370/ARMxy BL370 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL370/BLIIoT BL370 User Manual V1.0.docx',
    ],
    'BL410': [
        'ARM嵌入式控制器/BL410/ARMxy BL410 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL410/BLIIoT BL410 User Manual V1.0.docx',
    ],
    'BL440': [
        'ARM嵌入式控制器/BL440/ARMxy BL440 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL440/BLIIOT BL440 User Manual V1.0.docx',
    ],
    'BL450': [
        'ARM嵌入式控制器/BL450/ARMxy BL450 Datasheet V1.0.docx',
        'ARM嵌入式控制器/BL450/BLIIOT BL450 User Manual V1.0.docx',
    ],
    'BL460': [
        'ARM嵌入式控制器/BL460/ARMxy BL460 Datasheet V1.0.docx',
    ],
    # ===== ARM高性能网关 =====
    'BL116': [
        'ARM高性能网关/BL116/BLIIOT BL116 IIoT High-Performance Edge Gateway Data Sheet V1.0.docx',
        'ARM高性能网关/BL116/BLIIOT BL116 IIoT High-Performance Edge Gateway_User Manual_V1.0.docx',
    ],
    'BE116': [
        'ARM高性能网关/BE116/BLIIOT BE116 High-Performance Smart Grid Edge Gateways Data Sheet V1.0.docx',
        'ARM高性能网关/BE116/BLIIOT BE116 High-Performance Smart Grid Edge Gateways_User Manual_V1.0.docx',
    ],
    'BA116': [
        'ARM高性能网关/BA116/BLIIOT BA116 High-Performance Building HVAC Edge Gateway Data Sheet V1.0.docx',
        'ARM高性能网关/BA116/BLIIOT BA116 Series High-Performance Building HVAC Gateway_User Manual_V1.0.docx',
    ],
    'BL118': [
        'ARM高性能网关/BL118/BLIIOT BL118 Node-RED Edge Gateway Datasheet V1.0.docx',
        'ARM高性能网关/BL118/BLIIOT BL118 Node-RED Edge Gateway User Manual V1.3.docx',
    ],
    # ===== BL系列塑胶壳新型网关 =====
    'BL10_Series_Datasheets': ['BL系列塑胶壳新型网关/英文规格书/BL10系列规格书/' + f for f in os.listdir(base + '/BL系列塑胶壳新型网关/英文规格书/BL10系列规格书/') if f.endswith('.docx') and '$' not in f],
    'BA10_Series_Datasheets': ['BL系列塑胶壳新型网关/英文规格书/BA10系列规格书/' + f for f in os.listdir(base + '/BL系列塑胶壳新型网关/英文规格书/BA10系列规格书/') if f.endswith('.docx') and '$' not in f],
    'BE10_Series_Datasheets': ['BL系列塑胶壳新型网关/英文规格书/BE10系列规格书/' + f for f in os.listdir(base + '/BL系列塑胶壳新型网关/英文规格书/BE10系列规格书/') if f.endswith('.docx') and '$' not in f],
    'BL120_Series_Datasheets': ['BL系列塑胶壳新型网关/英文规格书/BL120系列规格书/' + f for f in os.listdir(base + '/BL系列塑胶壳新型网关/英文规格书/BL120系列规格书/') if f.endswith('.docx') and '$' not in f],
    'BL121_Series_Datasheets': ['BL系列塑胶壳新型网关/英文规格书/BL121系列规格书/' + f for f in os.listdir(base + '/BL系列塑胶壳新型网关/英文规格书/BL121系列规格书/') if f.endswith('.docx') and '$' not in f],
    'BE120_Datasheet': ['BL系列塑胶壳新型网关/英文规格书/BE120系列规格书/BLIIOT BE120 Data Sheet.docx'],
    # ===== R40 路由器 =====
    'R40': [
        '2025新版R40系列/规格书/BLIIOT R40 Industrial Cellular Router Data Sheet V2.0.docx',
        '2025新版R40系列/规格书/BLIIOT R40A R40B Cellular IIoT Router Data Sheet V2.0.docx',
    ],
    # ===== IOy系列 =====
    'IOy_Series': [
        'IOy系列多功能可编程远程IO/规格书/BLIIOT BA190 IOy Series BACnetIP Edge IO Module Data Sheet V1.0.docx',
        'IOy系列多功能可编程远程IO/规格书/BLIIOT BE190 IOy Series IEC104 Edge IO Module Data Sheet V1.0.docx',
        'IOy系列多功能可编程远程IO/规格书/BLIIOT BL190 IOy Series Modbus TCP Edge IO Module Data Sheet V1.0.docx',
        'IOy系列多功能可编程远程IO/规格书/BLIIOT BL191 IOy Series OPC UA Edge IO Module Data Sheet V1.0.docx',
        'IOy系列多功能可编程远程IO/规格书/BLIIOT BL192 IOy Series MQTT Edge IO Module Data Sheet V1.0.docx',
        'IOy系列多功能可编程远程IO/规格书/BLIIOT BL193 IOy Series SNMP Edge IO Module Data Sheet V1.0.docx',
    ],
    # ===== RTU50xx系列 =====
    'RTU5024': [
        'RTU50xx系列/RTU5024/产品规格书/RTU5024_Data Sheet_V1.7.docx',
        'RTU50xx系列/RTU5024/产品说明书/BLIIOT RTU5024_UserManual_V1.7.1.docx',
    ],
    'RTU5025': [
        'RTU50xx系列/RTU5025/产品规格书/RTU5025_Data Sheet_V1.9.docx',
        'RTU50xx系列/RTU5025/产品说明书/BLIIoT RTU5025 User Manual V2.1.docx',
    ],
    'RTU5028E': ['RTU50xx系列/RTU5028E/产品说明书/BLIIoT RTU5028E_usermanual_v1.1.docx'],
    'RTU5034': [
        'RTU50xx系列/RTU5034/产品规格书/RTU5034_Data Sheet_V1.0.docx',
        'RTU50xx系列/RTU5034/产品说明书/RTU5034_UserManual_V1.0.docx',
    ],
    # ===== S27x系列 =====
    'S271': ['S27x系列/S270_271/产品说明书/BLIIoT S271_User Manual_V1.0.docx'],
    'S275': ['S27x系列/S275/说明书/BLIIoT S275 _User Manual_V1.3.docx'],
    # ===== S150 =====
    'S150': ['S130-150/产品说明书/BLIIoT S150_User Manual_V3.0.1.docx'],
    # ===== IPM100 =====
    'IPM100': ['IPM100 IP67 IO模块/BLIIOT IPM100 IP67 Remote IO User Manual V1.0.docx'],
    # ===== 隔离器 =====
    'BL150_Isolators': ['隔离器/规格书/' + f for f in os.listdir(base + '/隔离器/规格书/') if f.endswith('.docx') and '$' not in f],
    # ===== BLIoTLink =====
    'BLIoTLink': ['BLIoTLink/BLIoTLink V1.0 User Manual.docx'],
    # ===== BLRAT =====
    'BLRAT': ['BLRAT/BLIIOT BLRAT User Manual.doc'],
    # ===== QuickConfig =====
    'QuickConfig': ['QuickConfig/BLIIOT QuickConfig User Manual_V1.0.docx'],
    # ===== 交换机 =====
    'BL16_Switches': ['交换机/BL16系列交换机/BL16系列工业以太网交换机产品规格书/' + f for f in os.listdir(base + '/交换机/BL16系列交换机/BL16系列工业以太网交换机产品规格书/') if f.endswith('.docx') and '$' not in f],
}

def extract_text(filepath):
    full_path = os.path.join(base, filepath)
    if not os.path.exists(full_path):
        return f"[FILE NOT FOUND: {full_path}]"
    try:
        doc = docx.Document(full_path)
        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n=== TABLE {ti+1} ===")
            for row in table.rows:
                cells = [c.text.strip()[:100] for c in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as e:
        return f"[ERROR: {e}]"

results = {}
for name, paths in key_files.items():
    print(f"\n=== READING: {name} ===")
    all_text = []
    for p in paths:
        text = extract_text(p)
        all_text.append(f"\n--- File: {p} ---\n{text}")
    results[name] = "\n".join(all_text)
    print(f"  -> {len(results[name])} chars")

# Save all extracted data
os.makedirs(os.path.dirname(r'C:\Users\Admin\AppData\Local\hermes\memories\脚本缓存\产品规格书\extracted_data.json'), exist_ok=True)
with open(r'C:\Users\Admin\AppData\Local\hermes\memories\脚本缓存\产品规格书\extracted_data.json', 'w', encoding='utf-8') as f:
    json.dump(results, f, ensure_ascii=False, indent=2)

print(f"\n\n✅ ALL EXTRACTED. Total products: {len(results)}")
print(f"Saved to: C:\\Users\\Admin\\AppData\\Local\\hermes\\memories\\脚本缓存\\产品规格书\\extracted_data.json")