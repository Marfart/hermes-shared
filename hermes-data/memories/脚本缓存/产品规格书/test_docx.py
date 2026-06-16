"""Test docx extraction."""
import docx

path = r'C:\Users\Admin\Desktop\Working\产品规格书\英文资料\ARM嵌入式控制器\BL310\ARMxy BL310 Datasheet V1.0.docx'
doc = docx.Document(path)
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
print(f'Sections: {len(doc.sections)}')

# Print first 30 paragraphs with content
count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f'P{i}: {t[:300]}')
        count += 1
    if count >= 20:
        break

# Print first 3 tables
for ti, table in enumerate(doc.tables[:3]):
    print(f'\n=== TABLE {ti+1} ===')
    for ri, row in enumerate(table.rows[:15]):
        cells = [c.text.strip()[:60] for c in row.cells]
        print(f'R{ri}: {" | ".join(cells)}')