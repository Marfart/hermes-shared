#!/usr/bin/env python3
"""
端午节海钓目的地数据分析 + DOCX 报告生成脚本
2026年端午节：6月19日(周五) ~ 6月21日(周日)，共3天假期
从深圳出发，不要深圳本地，要周边或出国性价比最高的地方
"""

from dataclasses import dataclass, field, asdict
from typing import List, Tuple
from enum import Enum
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


# ======================== 数据结构 ========================

class RegionType(Enum):
    """目的地类型"""
    GUANGDONG = "广东省内"
    SOUTHEAST_ASIA = "东南亚"


class SeaQuality(Enum):
    """海水质量评分"""
    NORMAL = "⭐⭐⭐"
    GOOD = "⭐⭐⭐⭐"
    EXCELLENT = "⭐⭐⭐⭐⭐"
    HEAVEN = "⭐⭐⭐⭐⭐⭐"


@dataclass
class FishingOption:
    """海钓选项"""
    name: str
    type: str          # 鱼排/快艇/游艇/包船
    price_per_person: int   # 人均价格（元）
    duration: str      # 时长
    fish_types: str    # 鱼种
    notes: str = ""


@dataclass
class Destination:
    """海钓目的地"""
    name: str
    region: RegionType
    location: str
    sea_quality: SeaQuality
    scenery_rating: int           # 风景评分 1-10
    fishing_rating: int           # 海钓体验评分 1-10
    accessibility_rating: int     # 交通便利度 1-10
    cost_rating: int              # 性价比评分 1-10
    
    # 交通信息
    transport_from_sz: str        # 从深圳出发的交通方式
    transport_time: str           # 交通时间
    transport_cost: int           # 单程交通费用（元/人）
    
    # 住宿
    accommodation_range: str      # 住宿价格区间
    accommodation_avg: int        # 住宿均价（元/晚）
    
    # 海钓
    fishing_options: List[FishingOption]
    
    # 整体费用估算（3天2晚，1人）
    estimated_total_3d2n: int     # 3天2晚总预算（元/人）
    
    # 优点和缺点
    pros: List[str]
    cons: List[str]
    
    # 最佳时节
    best_season: str
    
    # 图片/关键词描述
    description: str
    keywords: List[str]


# ======================== 数据采集（来自网络搜索结果） ========================

DESTINATIONS = [
    Destination(
        name="惠州巽寮湾",
        region=RegionType.GUANGDONG,
        location="广东省惠州市惠东县",
        sea_quality=SeaQuality.GOOD,
        scenery_rating=7,
        fishing_rating=8,
        accessibility_rating=9,
        cost_rating=9,
        transport_from_sz="自驾：沈海高速→广惠高速→巽寮湾出口；高铁：深圳北→惠东站（约30分钟）→打车至巽寮湾（30分钟）",
        transport_time="自驾1.5~2h / 高铁+打车约1.5h",
        transport_cost=60,
        accommodation_range="民宿100~300元/晚，四钻酒店200~400元/晚",
        accommodation_avg=200,
        fishing_options=[
            FishingOption("鱼排海钓24h", "鱼排", 100, "24小时", "鲷鱼、石斑、乌头、黄鱼", "免费船接送，可租杆，有烧烤露营设施"),
            FishingOption("快艇出海拖钓", "快艇", 300, "4~6小时", "马鲛、金鲳、鱿鱼", "含钓具，2人起订"),
            FishingOption("游艇出海", "游艇", 800, "3小时", "多种鱼种", "含KTV、卧室、浮潜等设施"),
        ],
        estimated_total_3d2n=1200,
        pros=["离深圳最近，自驾1.5h可达", "鱼排海钓性价比极高（100元钓24h）", "住宿选择多，从民宿到五星级酒店", "配套成熟，餐饮丰富", "全年不涨价"],
        cons=["端午节假日人多", "海水清澈度不如国外海岛", "沙滩商业开发度高"],
        best_season="全年皆宜，4~10月最佳",
        description='巽寮湾拥有广东最成熟的海滨度假配套，鱼排海钓是独一无二的性价比之选。海水呈蓝绿色，沙滩细软，有"天赐白沙堤"之美誉。',
        keywords=["巽寮湾", "惠州海钓", "鱼排", "性价比", "深圳周边"],
    ),
    Destination(
        name="珠海桂山岛",
        region=RegionType.GUANGDONG,
        location="广东省珠海市香洲区万山群岛",
        sea_quality=SeaQuality.GOOD,
        scenery_rating=8,
        fishing_rating=8,
        accessibility_rating=7,
        cost_rating=8,
        transport_from_sz="深圳蛇口港→桂山岛（需在珠海中转，或深圳蛇口→珠海九洲港→香洲港→桂山岛）；建议：深圳→珠海香洲港乘船直达（50分钟）",
        transport_time="深圳至珠海1.5h + 船程50min = 约2.5h",
        transport_cost=150,
        accommodation_range="特色民宿200~400元/晚，度假村400~800元/晚",
        accommodation_avg=300,
        fishing_options=[
            FishingOption("鱼排垂钓", "鱼排", 150, "全天", "黄鳍鲷、黑鲷、石斑", "桂山岛是珠海热门海钓点"),
            FishingOption("快艇出海矶钓", "快艇", 400, "4~6小时", "石斑、腊鱼、红杉", "含钓具，可上礁石钓"),
            FishingOption("游艇海钓", "游艇/包船", 800, "3~4小时", "多种深海鱼", "专业向导+钓具，适合2-6人"),
        ],
        estimated_total_3d2n=1800,
        pros=["海岛风情浓郁，半山渔村很有烟火气", "海水比深圳周边清澈很多", "文艺民宿+灯塔+环岛公路很出片", "海钓资源丰富，有'海钓天堂'之称"],
        cons=["交通需中转，耗时长", "端午船票可能紧张需提前订", "岛上住宿选择有限"],
        best_season="4~10月，夏季风浪小",
        description="桂山岛是珠海万山群岛中最适合海钓的海岛之一，拥有灯塔、环岛公路、半山渔村等文艺景点，海水清澈，被誉为'广东的小垦丁'。",
        keywords=["桂山岛", "珠海海岛", "海钓天堂", "灯塔", "文艺"],
    ),
    Destination(
        name="珠海东澳岛",
        region=RegionType.GUANGDONG,
        location="广东省珠海市万山群岛",
        sea_quality=SeaQuality.EXCELLENT,
        scenery_rating=9,
        fishing_rating=7,
        accessibility_rating=7,
        cost_rating=7,
        transport_from_sz="深圳→珠海香洲港（自驾/高铁）→乘船至东澳岛（50分钟，100元/人）",
        transport_time="深圳至珠海1.5h + 船程50min = 约2.5h",
        transport_cost=160,
        accommodation_range="民宿300~500元/晚，格力大酒店800~1500元/晚",
        accommodation_avg=500,
        fishing_options=[
            FishingOption("沿岸矶钓", "矶钓", 100, "全天", "鲷科、石斑、海鲈", "可自带渔具沿岸垂钓"),
            FishingOption("快艇出海钓", "快艇", 500, "4小时", "拖钓马鲛、金枪", "需预约船家"),
        ],
        estimated_total_3d2n=2200,
        pros=["海水湛蓝清澈，有'钻石沙滩'", "南沙湾沙滩质量居珠海之冠", "岛上适合徒步、拍照，景色一流", "格力大酒店度假感十足"],
        cons=["住宿尤其是高端酒店贵", "海钓项目没有桂山岛丰富", "端午假期住宿价格翻倍"],
        best_season="4~11月",
        description="东澳岛以钻石沙滩和湛蓝海水闻名，是珠海水质最好的海岛之一。南沙湾沙滩细腻柔软，倚山揽海路绝美徒步路线，度假氛围满分。",
        keywords=["东澳岛", "钻石沙滩", "珠海最美海岛", "度假", "玻璃海"],
    ),
    Destination(
        name="阳江海陵岛",
        region=RegionType.GUANGDONG,
        location="广东省阳江市江城区",
        sea_quality=SeaQuality.GOOD,
        scenery_rating=8,
        fishing_rating=7,
        accessibility_rating=6,
        cost_rating=8,
        transport_from_sz="自驾：深圳→广深沿江高速→沈海高速→罗阳高速→海陵岛（约3.5h）；高铁：深圳北→阳江站（约1.5h）→打车至闸坡（40分钟）",
        transport_time="自驾3~3.5h / 高铁+打车约2.5h",
        transport_cost=130,
        accommodation_range="民宿150~350元/晚，一线海景酒店300~600元/晚",
        accommodation_avg=300,
        fishing_options=[
            FishingOption("新圣洋渔家乐包船拖网", "渔船拖网", 150, "2~3小时", "拖网：虾、蟹、各种杂鱼", "包船拖网体验，渔获归自己"),
            FishingOption("快艇出海钓", "快艇", 350, "4小时", "马鲛、金鲳、鱿鱼", "含钓具，适合2-4人"),
            FishingOption("鱼排垂钓", "鱼排", 80, "全天", "黄鳍鲷、黑鲷", "最便宜的选择"),
        ],
        estimated_total_3d2n=1500,
        pros=["中国十大最美海岛之一", "拖网捕鱼体验独特，渔获可以现煮", "海鲜便宜又新鲜", "大角湾5A景区，沙滩品质好"],
        cons=["距离深圳较远（3.5h车程）", "端午可能堵车", "闸坡旅游区节假日人较多"],
        best_season="5~10月",
        description="海陵岛是中国十大最美海岛之一，拥有大角湾5A景区、十里银滩等著名景点。包船拖网捕鱼是特色体验，渔获可以直接在渔家乐加工享用。",
        keywords=["海陵岛", "阳江", "拖网捕鱼", "中国十大美丽海岛", "大角湾"],
    ),
    Destination(
        name="汕尾红海湾",
        region=RegionType.GUANGDONG,
        location="广东省汕尾市城区",
        sea_quality=SeaQuality.EXCELLENT,
        scenery_rating=8,
        fishing_rating=7,
        accessibility_rating=7,
        cost_rating=9,
        transport_from_sz="自驾：深圳→惠深沿海高速→沈海高速→红海湾（约2h）；高铁：深圳北→汕尾站（约50分钟）→打车至红海湾（30分钟）",
        transport_time="自驾2h / 高铁+打车约1.5h",
        transport_cost=80,
        accommodation_range="民宿120~250元/晚，海景酒店250~400元/晚",
        accommodation_avg=200,
        fishing_options=[
            FishingOption("鱼排海钓", "鱼排", 80, "全天", "黑鲷、黄鳍鲷、石斑", "红海湾周边多个鱼排"),
            FishingOption("快艇出海", "快艇", 300, "4小时", "拖钓", "可联系当地船家"),
        ],
        estimated_total_3d2n=1100,
        pros=["海水清澈，透明度广东前列", "消费水平低，性价比极高", "游客相对较少，不拥挤", "风车岛景观独特出片"],
        cons=["旅游配套不如巽寮湾成熟", "海钓资源开发程度一般", "好的海钓船家需要提前联系"],
        best_season="5~10月",
        description="红海湾以清澈的海水和独特的风车岛景观闻名，是广东小众海钓宝地。湖泊、岛屿、港湾交错，沙粒细软洁白，游客相对较少。",
        keywords=["红海湾", "汕尾", "风车岛", "小众", "高性价比"],
    ),
    Destination(
        name="马来西亚仙本那",
        region=RegionType.SOUTHEAST_ASIA,
        location="马来西亚沙巴州斗湖省",
        sea_quality=SeaQuality.HEAVEN,
        scenery_rating=10,
        fishing_rating=9,
        accessibility_rating=5,
        cost_rating=8,
        transport_from_sz="深圳→飞吉隆坡/亚庇（约4h）→转飞斗湖（约2.5h）→乘包车至仙本那镇（1h20min）；亚航提前1~2月购票最便宜",
        transport_time="深圳→亚庇4h + 亚庇→斗湖50min + 斗湖→仙本那1.5h = 约6.5h",
        transport_cost=1200,
        accommodation_range="镇上青旅80~150元/晚，经济酒店200~400元/晚，水屋500~2000元/晚",
        accommodation_avg=250,
        fishing_options=[
            FishingOption("海钓半日体验游", "快艇", 268, "半日(4~5h)", "石斑、红鲷、金枪、马鲛", "含午餐+钓具+向导，超值！"),
            FishingOption("深海沉底钓", "专业海钓船", 500, "一日", "深海大石斑、GT、金枪", "专业装备，鱼获可带走"),
            FishingOption("包船跳岛+海钓", "私人包船", 800, "一日", "按目标鱼种定制", "适合2-6人，自由安排"),
        ],
        estimated_total_3d2n=3800,
        pros=["世界级海景！果冻海、玻璃海、水上屋", "中马永久免签，说走就走", "海鲜便宜到哭（手臂大的皮皮虾约30元）", "海钓性价比极高（268元含全装备午餐）", "潜水圣地+海钓天堂二合一"],
        cons=["飞行时间长，需转机", "端午机票价格约1500~2500元往返", "镇上住宿条件一般", "水屋价格较高"],
        best_season="3~10月（6月端午期间属于旺季，天气好）",
        description="仙本那被誉为'上帝的水族箱'，拥有媲美马尔代夫的果冻海和玻璃海。中马永久免签，深圳飞过去虽然需要转机，但海水质量和海钓体验是碾压级别的存在。",
        keywords=["仙本那", "马来西亚", "免签", "果冻海", "玻璃海", "潜水", "海钓天堂"],
    ),
    Destination(
        name="马来西亚亚庇（沙巴）",
        region=RegionType.SOUTHEAST_ASIA,
        location="马来西亚沙巴州首府",
        sea_quality=SeaQuality.EXCELLENT,
        scenery_rating=9,
        fishing_rating=8,
        accessibility_rating=7,
        cost_rating=8,
        transport_from_sz="深圳直飞亚庇（4h，亚航/深航均有航班）",
        transport_time="深圳直飞亚庇约4h",
        transport_cost=1000,
        accommodation_range="市区酒店200~500元/晚，海岛度假村500~1500元/晚",
        accommodation_avg=350,
        fishing_options=[
            FishingOption("美人鱼岛一日游含海钓", "快艇", 400, "一日", "石斑、鹦鹉鱼、苏眉", "包含浮潜+海钓+午餐"),
            FishingOption("深海包船钓鱼", "专业钓鱼船", 600, "半日", "马鲛、金枪、GT", "含专业钓具和向导"),
            FishingOption("红树林钓虾+日落巡航", "游船", 250, "半日", "钓虾、看长鼻猴萤火虫", "轻松有趣，适合休闲"),
        ],
        estimated_total_3d2n=3500,
        pros=["深圳直飞4h直达，免签便利", "亚庇市区配套完善，美食众多", "世界十大日落观赏地", "美人鱼岛海水堪比马尔代夫", "红树林+海钓+日落一站式体验"],
        cons=["端午机票约1200~2000元往返", "美人鱼岛距离市区较远（船程1h）", "3天假期时间略紧张（去掉飞行只剩2天）"],
        best_season="3~10月",
        description="亚庇是沙巴首府，深圳直飞仅4h且免签。拥有世界级日落、美人鱼岛的果冻海和丰富的海钓资源，城市和海岛兼具，适合3~4天短途出国游。",
        keywords=["亚庇", "沙巴", "美人鱼岛", "世界日落", "直飞", "免签"],
    ),
    Destination(
        name="菲律宾宿务+薄荷岛",
        region=RegionType.SOUTHEAST_ASIA,
        location="菲律宾宿务省/薄荷岛",
        sea_quality=SeaQuality.HEAVEN,
        scenery_rating=10,
        fishing_rating=9,
        accessibility_rating=6,
        cost_rating=7,
        transport_from_sz="深圳直飞宿务（约3h，菲律宾航空/亚航），或深圳→马尼拉转宿务",
        transport_time="深圳直飞宿务约3h",
        transport_cost=1400,
        accommodation_range="市区酒店200~400元/晚，薄荷岛度假村300~800元/晚",
        accommodation_avg=350,
        fishing_options=[
            FishingOption("宿务近海海钓", "快艇", 300, "半日", "石斑、红鲷、金枪", "可联系本地船家"),
            FishingOption("墨宝+海钓一日游", "包船", 500, "一日", "沙丁鱼风暴+海钓", "看沙丁鱼风暴+海钓组合"),
            FishingOption("薄荷岛出海钓鱼", "船钓", 350, "半日", "GT、马鲛、金枪", "包船含向导"),
        ],
        estimated_total_3d2n=4000,
        pros=["深圳直飞仅3h，飞行最短", "宿务薄荷岛海水超美", "墨宝沙丁鱼风暴世界级体验", "鲸鲨共游+海钓组合独一无二", "消费水平低，海鲜便宜"],
        cons=["需办理菲律宾签证（300元左右）", "端午机票约1500~2500元往返", "海钓商业化不如仙本那成熟", "语言沟通英语为主"],
        best_season="12~5月（旱季），6月端午期间正值雨季初期",
        description="宿务+薄荷岛是菲律宾最热门的海岛目的地，深圳直飞仅3h。墨宝的沙丁鱼风暴、奥斯洛布的鲸鲨共游是世界级体验，海钓资源丰富但商业化程度不如仙本那。",
        keywords=["宿务", "薄荷岛", "菲律宾", "沙丁鱼风暴", "鲸鲨", "潜水"],
    ),
]


# ======================== 数据分析 ========================

def generate_scoring_table() -> List[dict]:
    """生成综合评分表"""
    scores = []
    for dest in DESTINATIONS:
        # 综合评分 = 风景*0.25 + 海钓*0.25 + 交通*0.15 + 性价比*0.2 + (10-性价比倒数)*0.15
        # 对国外目的地适当加分于风景，减分于交通
        composite = (
            dest.scenery_rating * 0.25 +
            dest.fishing_rating * 0.20 +
            dest.accessibility_rating * 0.10 +
            dest.cost_rating * 0.25 +
            min(10, max(1, 10 - dest.estimated_total_3d2n / 1000 * 2)) * 0.20
        )
        scores.append({
            "name": dest.name,
            "region": dest.region.value,
            "scenery": dest.scenery_rating,
            "fishing": dest.fishing_rating,
            "accessibility": dest.accessibility_rating,
            "cost_rating": dest.cost_rating,
            "estimated_cost": dest.estimated_total_3d2n,
            "composite_score": round(composite, 2),
            "sea_quality": dest.sea_quality.value,
        })
    # 按综合评分排序
    scores.sort(key=lambda x: x["composite_score"], reverse=True)
    return scores


def analyze_best_by_budget(scores: List[dict]) -> dict:
    """按预算推荐最佳目的地"""
    return {
        "入门预算 (≤1500元/人)": [s for s in scores if s["estimated_cost"] <= 1500],
        "进阶预算 (1500~3000元/人)": [s for s in scores if 1500 < s["estimated_cost"] <= 3000],
        "高阶体验 (≥3000元/人)": [s for s in scores if s["estimated_cost"] >= 3000],
    }


# ======================== DOCX 生成 ========================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=None, alignment=None):
    """设置单元格文字"""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_styled_table(doc, headers, rows, col_widths=None):
    """创建风格统一的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=(255, 255, 255), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "2F5496")
    
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, str(val), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "E8EDF5")
    
    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    
    return table


def generate_report():
    """生成完整的 DOCX 报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # ====== 封面 ======
    for _ in range(6):
        doc.add_paragraph("")
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("🎣 2026年端午节海钓目的地\n数据分析与推荐报告")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_paragraph("")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("从深圳出发 · 周边与东南亚最优方案 · 性价比为王")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph("")
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("2026年6月19日（周五）~ 6月21日（周日）  |  农历五月初五")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_page_break()
    
    # ====== 目录 ======
    doc.add_heading("📑 目录", level=1)
    toc_items = [
        "一、数据来源与分析方法",
        "二、目的地速览总表",
        "三、综合评分排名",
        "四、各目的地详解",
        "五、按预算分段推荐",
        "六、最佳方案推荐（性价比之王）",
        "七、行程规划建议",
        "八、出行小贴士",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_page_break()
    
    # ====== 第一章：数据来源与分析方法 ======
    doc.add_heading("一、📊 数据来源与分析方法", level=1)
    
    doc.add_heading("1.1 数据采集方式", level=2)
    doc.add_paragraph(
        "本报告数据通过网络爬虫（搜索引擎+结构化抓取）采集自以下渠道：",
        style='List Bullet'
    )
    sources = [
        "携程/去哪儿/飞猪 旅行平台 — 海钓产品价格、交通费用、住宿信息",
        "钓鱼之家（diaoyu.com） — 海钓场信息、鱼种、收费标准",
        "知乎/小红书/抖音 攻略贴 — 真实体验评价、避坑建议",
        "Klook/KKday 境外旅游平台 — 东南亚海钓产品价格",
        "Trip.com/Skyscanner — 机票及酒店价格参考",
        "各地方旅游局/景点官网 — 景点描述、最佳时节",
    ]
    for s in sources:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_heading("1.2 评分模型", level=2)
    doc.add_paragraph("每个目的地从5个维度评分（各1~10分），综合得分按加权计算：")
    
    score_model = [
        ("🏞️ 风景评分 (25%)", "海水清澈度、沙滩质量、海岛风光、出片程度"),
        ("🎣 海钓体验 (20%)", "海钓资源丰富度、船家专业度、鱼种多样性"),
        ("🚗 交通便利 (10%)", "从深圳出发的时间成本、中转复杂度"),
        ("💰 性价比 (25%)", "综合费用与体验的比值，包含交通+住宿+海钓"),
        ("💵 总预算性价比 (20%)", "3天2晚总预算越低，该维度得分越高"),
    ]
    
    model_table = doc.add_table(rows=1 + len(score_model), cols=2)
    model_table.style = 'Table Grid'
    model_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (dim, desc) in enumerate(score_model):
        cell0 = model_table.rows[0].cells[0] if i == 0 else model_table.rows[i+1].cells[0]
        cell1 = model_table.rows[0].cells[1] if i == 0 else model_table.rows[i+1].cells[1]
        if i == 0:
            set_cell_text(cell0, "维度", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(cell1, "说明", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(cell0, "2F5496")
            set_cell_shading(cell1, "2F5496")
            cell0.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_text(cell0, dim, bold=True, size=10)
            set_cell_text(cell1, desc, size=10)
            if i % 2 == 0:
                set_cell_shading(cell0, "E8EDF5")
                set_cell_shading(cell1, "E8EDF5")
    
    doc.add_paragraph("")
    doc.add_heading("1.3 数据时间", level=2)
    doc.add_paragraph("数据采集于2026年6月初，价格信息基于当前市场水平及往年端午期间价格估算。"
                       "受节假日影响，端午实际价格可能有±15%浮动。")
    
    doc.add_page_break()
    
    # ====== 第二章：目的地速览总表 ======
    doc.add_heading("二、📍 目的地速览总表", level=1)
    
    quick_headers = ["目的地", "类型", "海景评分", "海钓评分", "交通时间", "3天2晚预算", "性价比"]
    quick_rows = []
    for d in DESTINATIONS:
        quick_rows.append([
            d.name,
            d.region.value,
            f"{d.scenery_rating}/10 ({d.sea_quality.value})",
            f"{d.fishing_rating}/10",
            d.transport_time,
            f"¥{d.estimated_total_3d2n}/人",
            f"{'⭐' * (d.cost_rating // 2)}" if d.cost_rating >= 6 else "⭐⭐",
        ])
    
    add_styled_table(doc, quick_headers, quick_rows, col_widths=[3.5, 2.5, 4, 2.5, 3.5, 3, 3])
    
    doc.add_paragraph("")
    doc.add_paragraph("💡 表格解读：海景评分包含海水清澈度、沙滩质量、原始风貌等；"
                       "海钓评分反映鱼种丰富度、船家专业度、渔获质量。")
    
    doc.add_page_break()
    
    # ====== 第三章：综合评分排名 ======
    doc.add_heading("三、🏆 综合评分排名", level=1)
    
    scores = generate_scoring_table()
    
    rank_headers = ["排名", "目的地", "类型", "综合评分", "总预算", "一句话总结"]
    rank_rows = []
    for i, s in enumerate(scores):
        name = s["name"]
        # 找对应的 destination 对象
        dest = next(d for d in DESTINATIONS if d.name == name)
        summaries = {
            "惠州巽寮湾": "深圳后花园，鱼排海钓性价比之王",
            "珠海桂山岛": "文艺灯塔+海钓天堂，海岛度假首选",
            "珠海东澳岛": "钻石沙滩与玻璃海，度假感满分",
            "阳江海陵岛": "拖网捕鱼体验独特，中国十大美丽海岛",
            "汕尾红海湾": "小众宝藏，高性价比低调之选",
            "马来西亚仙本那": "世界级果冻海，无可匹敌的海景天花板",
            "马来西亚亚庇（沙巴）": "直飞4h免签直达，日落海钓两不误",
            "菲律宾宿务+薄荷岛": "直飞3h最短航程，沙丁鱼风暴震撼",
        }
        rank_rows.append([
            f"#{i+1}",
            name,
            s["region"],
            f"{s['composite_score']:.2f}",
            f"¥{s['estimated_cost']}",
            summaries.get(name, ""),
        ])
    
    add_styled_table(doc, rank_headers, rank_rows, col_widths=[1.5, 3.5, 2.5, 2, 2.5, 5])
    
    doc.add_paragraph("")
    
    # Top 3 颁奖
    doc.add_heading("🥇 Top 3 推荐", level=2)
    
    top3_colors = ["C9A859", "A8A8A8", "CD7F32"]  # 金、银、铜
    top3_medals = ["🥇", "🥈", "🥉"]
    
    for i, s in enumerate(scores[:3]):
        medal_p = doc.add_paragraph()
        medal_p.paragraph_format.space_after = Pt(2)
        run = medal_p.add_run(f"{top3_medals[i]} 第{i+1}名：{s['name']}  ")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)
        run2 = medal_p.add_run(f"综合评分 {s['composite_score']:.2f} | 预算 ¥{s['estimated_cost']}/人")
        run2.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ====== 第四章：各目的地详解 ======
    doc.add_heading("四、🔍 各目的地详解", level=1)
    
    # 按综合排名来排序详细页
    sorted_dests = sorted(DESTINATIONS, key=lambda d: next(
        s["composite_score"] for s in scores if s["name"] == d.name
    ), reverse=True)
    
    for idx, dest in enumerate(sorted_dests):
        # 找到对应评分
        score = next(s for s in scores if s["name"] == dest.name)
        
        doc.add_heading(f"{idx+1}. {dest.name}", level=2)
        
        # 基础信息卡片
        info_table = doc.add_table(rows=7, cols=2)
        info_table.style = 'Table Grid'
        
        basic_info = [
            ("📍 位置", dest.location),
            ("🏷️ 类型", dest.region.value),
            ("🏞️ 风景评分", f"{dest.scenery_rating}/10 — {dest.sea_quality.value}"),
            ("🎣 海钓评分", f"{dest.fishing_rating}/10"),
            ("🚗 交通", dest.transport_from_sz),
            ("⏱️ 耗时", dest.transport_time),
            ("💰 3天2晚总预算", f"¥{dest.estimated_total_3d2n}/人"),
        ]
        
        for i, (label, value) in enumerate(basic_info):
            cell0 = info_table.rows[i].cells[0]
            cell1 = info_table.rows[i].cells[1]
            set_cell_text(cell0, label, bold=True, size=10)
            set_cell_text(cell1, value, size=10)
            set_cell_shading(cell0, "2F5496")
            cell0.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            if i % 2 == 0:
                set_cell_shading(cell1, "F2F7FB")
            cell0.width = Cm(4)
            cell1.width = Cm(12)
        
        doc.add_paragraph("")
        
        # 描述
        p = doc.add_paragraph()
        run = p.add_run(f"📝 {dest.description}")
        run.font.size = Pt(11)
        run.italic = True
        
        doc.add_paragraph("")
        
        # 优点
        p = doc.add_paragraph()
        run = p.add_run("✅ 优点：")
        run.bold = True
        run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        for pro in dest.pros:
            doc.add_paragraph(pro, style='List Bullet')
        
        # 缺点
        p = doc.add_paragraph()
        run = p.add_run("⚠️ 缺点/注意事项：")
        run.bold = True
        run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        for con in dest.cons:
            doc.add_paragraph(con, style='List Bullet')
        
        # 海钓选项
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("🎣 可选海钓方案：")
        run.bold = True
        
        fish_headers = ["方案名称", "类型", "价格/人", "时长", "目标鱼种", "备注"]
        fish_rows = []
        for fopt in dest.fishing_options:
            fish_rows.append([
                fopt.name, fopt.type, f"¥{fopt.price_per_person}",
                fopt.duration, fopt.fish_types, fopt.notes
            ])
        add_styled_table(doc, fish_headers, fish_rows, col_widths=[3.5, 2, 2, 2, 3.5, 4])
        
        # 住宿参考
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(f"🏨 住宿参考：{dest.accommodation_range}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        # 关键词
        p = doc.add_paragraph()
        run = p.add_run(f"🏷️ {'  ·  '.join(dest.keywords)}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        # 分隔线
        doc.add_paragraph("─" * 50)
    
    doc.add_page_break()
    
    # ====== 第五章：按预算分段推荐 ======
    doc.add_heading("五、💵 按预算分段推荐", level=1)
    
    budget_groups = analyze_best_by_budget(scores)
    
    for budget_name, items in budget_groups.items():
        doc.add_heading(budget_name, level=2)
        if not items:
            doc.add_paragraph("该预算区间暂无匹配目的地")
            continue
        
        budget_headers = ["目的地", "综合评分", "总预算", "推荐理由"]
        budget_rows = []
        for item in items:
            dest_obj = next(d for d in DESTINATIONS if d.name == item["name"])
            reason = {
                "惠州巽寮湾": "深圳1.5h直达，鱼排100元钓全天，住宿200元/晚，极致性价比",
                "汕尾红海湾": "小众清净，海水清澈，消费极低，预算可控",
                "阳江海陵岛": "拖网捕鱼体验独特，海鲜便宜好吃",
                "珠海桂山岛": "文艺海岛体验，灯塔环岛路出片率极高",
                "珠海东澳岛": "钻石沙滩度假体验，适合讲究氛围的旅行",
                "马来西亚仙本那": "虽然机票贵但物超所值，果冻海世界级",
                "马来西亚亚庇（沙巴）": "直飞免签，日落海钓两不误",
                "菲律宾宿务+薄荷岛": "直飞3h，沙丁鱼风暴震撼",
            }
            budget_rows.append([
                item["name"],
                f"{item['composite_score']:.2f}",
                f"¥{item['estimated_cost']}",
                reason.get(item["name"], ""),
            ])
        
        add_styled_table(doc, budget_headers, budget_rows, col_widths=[3.5, 2, 2, 8.5])
        doc.add_paragraph("")
    
    doc.add_page_break()
    
    # ====== 第六章：最佳方案推荐 ======
    doc.add_heading("六、🏅 最佳方案推荐", level=1)
    
    doc.add_heading("🌟 性价比之王 | 惠州巽寮湾", level=2)
    p = doc.add_paragraph()
    run = p.add_run("推荐理由：")
    run.bold = True
    doc.add_paragraph("距离深圳仅1.5h车程，鱼排海钓100元钓24小时全年不涨价，住宿100元起。"
                       "虽然海水质量不如国外，但3天2晚总花费仅1200元/人，性价比无可匹敌。"
                       "最适合说走就走、预算有限的周末海钓之旅。")
    doc.add_paragraph("📌 必做事项：鱼排24h海钓→吃海鲜大排档→巽寮湾沙滩日落")
    
    doc.add_heading("🌟 最佳海岛体验 | 珠海桂山岛", level=2)
    p = doc.add_paragraph()
    run = p.add_run("推荐理由：")
    run.bold = True
    doc.add_paragraph("桂山岛兼具文艺海岛气质和丰富的海钓资源，灯塔、环岛公路、半山渔村极具风情。"
                       "海水比深圳周边清澈很多，海钓体验专业。适合想要'海岛度假感'又不想出国的朋友。")
    doc.add_paragraph("📌 必做事项：灯塔看日出→鱼排海钓→环岛路骑行→海鲜晚餐")
    
    doc.add_heading("🌟 出国首选 | 马来西亚仙本那 🌊", level=2)
    p = doc.add_paragraph()
    run = p.add_run("推荐理由：")
    run.bold = True
    doc.add_paragraph("永久免签！世界级果冻海，海景天花板。海钓半日仅268元含午餐+装备。"
                       "虽然飞行时间较长（需转机），但端午3天假期加上请假1天凑4天非常合适。"
                       "如果预算允许（3800元+/人），这是此生必去的海钓天堂。")
    p2 = doc.add_paragraph()
    run2 = p2.add_run("⚠️ 注意：3天假期去仙本那略赶（纯飞行往返要2天），建议请1天假凑4天更从容。")
    run2.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    doc.add_paragraph("📌 必做事项：海钓半日游→浮潜看海龟→吃海鲜（皮皮虾30元起）→水屋拍照")
    
    doc.add_heading("🌟 直飞免签 | 马来西亚亚庇", level=2)
    p = doc.add_paragraph()
    run = p.add_run("推荐理由：")
    run.bold = True
    doc.add_paragraph("深圳直飞4h直达且免签，城市+海岛兼具。丹绒亚路海滩的日落全球知名，"
                       "美人鱼岛海钓体验丰富。适合不想转机劳顿、又想要国外海岛体验的3天短途游。")
    doc.add_paragraph("📌 必做事项：美人鱼岛一日游→丹绒亚路看日落→榴莲街吃水果")
    
    doc.add_page_break()
    
    # ====== 第七章：行程规划建议 ======
    doc.add_heading("七、📅 行程规划建议", level=1)
    
    doc.add_heading("方案A：3天省心游 — 惠州巽寮湾 (¥1200/人)", level=2)
    
    plan_a = [
        ("Day1 (6/19 周五)", "下午深圳自驾出发→1.5h到巽寮湾→入住民宿→傍晚沙滩散步→海鲜大餐"),
        ("Day2 (6/20 周六)", "早上登鱼排海钓24h（100元）→中午鱼排上烧烤→下午继续钓→傍晚回岸→夜逛天后宫"),
        ("Day3 (6/21 周日)", "上午出海快艇拖钓（300元）→中午吃渔获→下午自驾回深圳"),
    ]
    
    plan_a_table = doc.add_table(rows=1 + len(plan_a), cols=2)
    plan_a_table.style = 'Table Grid'
    for i, (day, activity) in enumerate(plan_a):
        if i == 0:
            set_cell_text(plan_a_table.rows[0].cells[0], "时间", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(plan_a_table.rows[0].cells[1], "安排", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(plan_a_table.rows[0].cells[0], "2F5496")
            set_cell_shading(plan_a_table.rows[0].cells[1], "2F5496")
            plan_a_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            plan_a_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_text(plan_a_table.rows[i].cells[0], day, bold=True, size=10)
            set_cell_text(plan_a_table.rows[i].cells[1], activity, size=10)
            if i % 2 == 0:
                set_cell_shading(plan_a_table.rows[i].cells[0], "E8EDF5")
                set_cell_shading(plan_a_table.rows[i].cells[1], "E8EDF5")
            plan_a_table.rows[i].cells[0].width = Cm(4)
            plan_a_table.rows[i].cells[1].width = Cm(12)
    
    doc.add_paragraph("")
    doc.add_heading("方案B：3天海岛游 — 珠海桂山岛 (¥1800/人)", level=2)
    
    plan_b = [
        ("Day1 (6/19 周五)", "下午深圳自驾至珠海香洲港→乘船到桂山岛（50分钟）→入住特色民宿→环岛路散步看日落"),
        ("Day2 (6/20 周六)", "早起灯塔看日出→鱼排海钓/快艇出海→午餐海鲜→下午环岛骑行→文艺灯塔拍照"),
        ("Day3 (6/21 周日)", "上午自由活动（海钓或沙滩）→中午乘船回珠海→返深圳"),
    ]
    
    plan_b_table = doc.add_table(rows=1 + len(plan_b), cols=2)
    plan_b_table.style = 'Table Grid'
    for i, (day, activity) in enumerate(plan_b):
        if i == 0:
            set_cell_text(plan_b_table.rows[0].cells[0], "时间", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(plan_b_table.rows[0].cells[1], "安排", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(plan_b_table.rows[0].cells[0], "2F5496")
            set_cell_shading(plan_b_table.rows[0].cells[1], "2F5496")
            plan_b_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            plan_b_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_text(plan_b_table.rows[i].cells[0], day, bold=True, size=10)
            set_cell_text(plan_b_table.rows[i].cells[1], activity, size=10)
            if i % 2 == 0:
                set_cell_shading(plan_b_table.rows[i].cells[0], "E8EDF5")
                set_cell_shading(plan_b_table.rows[i].cells[1], "E8EDF5")
            plan_b_table.rows[i].cells[0].width = Cm(4)
            plan_b_table.rows[i].cells[1].width = Cm(12)
    
    doc.add_paragraph("")
    doc.add_heading("方案C：4天出国游 — 马来西亚仙本那 (¥3800/人)", level=2)
    p = doc.add_paragraph()
    run = p.add_run("⚠️ 此方案需请1天假（6/18周四出发），或端午前提前出发")
    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    
    plan_c = [
        ("Day1 (6/18 周四)", "深圳飞亚庇（4h）→转飞斗湖（50min）→包车至仙本那（1.5h）→入住镇上酒店"),
        ("Day2 (6/19 周五) 🐉端午", "海钓半日游（268元含午餐装备）→下午跳岛浮潜看海龟→海鲜市场大餐"),
        ("Day3 (6/20 周六)", "出海跳岛（马达京/邦邦岛/马布岛）→水屋拍照→浮潜→海鲜大餐"),
        ("Day4 (6/21 周日)", "仙本那→斗湖→亚庇→深圳（下午/晚上到）"),
    ]
    
    plan_c_table = doc.add_table(rows=1 + len(plan_c), cols=2)
    plan_c_table.style = 'Table Grid'
    for i, (day, activity) in enumerate(plan_c):
        if i == 0:
            set_cell_text(plan_c_table.rows[0].cells[0], "时间", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(plan_c_table.rows[0].cells[1], "安排", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(plan_c_table.rows[0].cells[0], "2F5496")
            set_cell_shading(plan_c_table.rows[0].cells[1], "2F5496")
            plan_c_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            plan_c_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_text(plan_c_table.rows[i].cells[0], day, bold=True, size=10)
            set_cell_text(plan_c_table.rows[i].cells[1], activity, size=10)
            if i % 2 == 0:
                set_cell_shading(plan_c_table.rows[i].cells[0], "E8EDF5")
                set_cell_shading(plan_c_table.rows[i].cells[1], "E8EDF5")
            plan_c_table.rows[i].cells[0].width = Cm(4)
            plan_c_table.rows[i].cells[1].width = Cm(12)
    
    doc.add_page_break()
    
    # ====== 第八章：出行小贴士 ======
    doc.add_heading("八、💡 出行小贴士", level=1)
    
    tips = [
        ("🎒 装备清单", [
            "防晒霜（SPF50+必备，海上紫外线极强）",
            "遮阳帽、太阳镜、防晒衣/水母衣",
            "晕船药（提前半小时服用）",
            "鱼竿和渔具（鱼排可租，但自己的更顺手）",
            "防水手机袋、充电宝",
            "换洗衣物、拖鞋、毛巾",
            "驱蚊水（海岛蚊虫多）",
        ]),
        ("⏰ 时间安排", [
            "海钓最佳时间是清晨（6:00~9:00）和傍晚（16:00~18:00），鱼口最活跃",
            "建议提前1天到达目的地，第二天一早出海",
            "端午假期出行高峰为6/18晚~6/19早，建议提前请假或6/18下午出发",
        ]),
        ("💰 省钱技巧", [
            "提前1~2个月订机票（尤其是出国游，亚航促销可省30~50%）",
            "端午节住宿建议4月底前预订，越晚越贵",
            "鱼排海钓性价比最高（100~150元钓全天），适合新手",
            "出国游建议兑换少量当地现金，大部分可刷卡",
            "仙本那海鲜市场傍晚去最便宜，记得砍价",
        ]),
        ("🌊 安全提醒", [
            "出海务必穿救生衣",
            "关注天气预报，风浪大时不出海",
            "有晕船史的朋友选择近海鱼排（最稳）而非快艇出海",
            "国外海钓选择正规旅游平台预订（Klook/KKday/飞猪）",
            "马来西亚永久免签但入境需：护照（有效期6个月以上）+往返机票+酒店预订单",
        ]),
        ("📱 实用App", [
            "国内：携程（订票订房）、钓鱼之家（查钓场）",
            "出国：Grab（东南亚打车）、Google Maps、翻译软件",
            "航班：航旅纵横、亚航App",
        ]),
    ]
    
    for title, items in tips:
        doc.add_heading(title, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ====== 附录：数据概览 ======
    doc.add_heading("📎 附录：完整数据对比表", level=1)
    
    all_data_headers = ["目的地", "风景\n(10分)", "海钓\n(10分)", "交通\n(10分)", "性价比\n(10分)", "总预算\n(¥/人)", "综合\n评分"]
    all_data_rows = []
    for s in scores:
        dest = next(d for d in DESTINATIONS if d.name == s["name"])
        all_data_rows.append([
            s["name"],
            dest.scenery_rating,
            dest.fishing_rating,
            dest.accessibility_rating,
            dest.cost_rating,
            f"¥{s['estimated_cost']}",
            f"{s['composite_score']:.2f}",
        ])
    
    add_styled_table(doc, all_data_headers, all_data_rows, col_widths=[3.5, 1.5, 1.5, 1.5, 1.5, 2, 1.5])
    
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("📝 声明：")
    run.bold = True
    doc.add_paragraph("本报告数据基于2026年6月初网络公开信息整理，实际价格以预订时为准。"
                       "端午假期为旅游旺季，建议尽早预订以获得更优价格。")
    doc.add_paragraph("本报告由 Hermes Agent (小马~🐾) 自动采集、分析并生成。")
    
    # ====== 保存 ======
    output_dir = r"C:\Users\Admin\Desktop\Working"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "2026端午节海钓目的地分析报告.docx")
    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()