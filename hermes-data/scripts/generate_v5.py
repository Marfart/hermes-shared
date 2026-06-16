#!/c/Users/Admin/AppData/Local/Programs/Python/Python313/python.exe
# -*- coding: utf-8 -*-
"""v5 - Corrected: antenna 8dBi, battery 11.1V, IoT lock detail, payment terms,
   SmartGrid Africa, 50W solar panel $9.2, transparent tables"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree as ET
import copy

LOGO = r"C:\Users\Admin\Desktop\Working\津巴布尔客户\签署文件\钡铼英文LOGO.png"
OUT = r"C:\Users\Admin\AppData\Local\hermes\更新后新文件.docx"

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)

# ---- HELPERS ----

def mg(c, t=80, b=80, l=150, r=150):
    tcPr = c._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcMar')): tcPr.remove(e)
    tcPr.append(parse_xml('<w:tcMar %s><w:top w:w="%d" w:type="dxa"/><w:bottom w:w="%d" w:type="dxa"/><w:left w:w="%d" w:type="dxa"/><w:right w:w="%d" w:type="dxa"/></w:tcMar>' % (nsdecls("w"), t, b, l, r)))

def rf(run, nm="Calibri", ea="微软雅黑"):
    rPr = run._r.get_or_add_rPr()
    for e in rPr.findall(qn('w:rFonts')): rPr.remove(e)
    rPr.append(parse_xml('<w:rFonts %s w:ascii="%s" w:hAnsi="%s" w:eastAsia="%s"/>' % (nsdecls("w"), nm, nm, ea)))

def cll(c, txt, b=False, sz=10, clr="1E293B", al=None, lh=1.1):
    p = c.paragraphs[0]; p.clear()
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = lh
    if al: p.alignment = al
    r = p.add_run(txt); r.bold = b; r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor.from_string(clr); rf(r)

def tbl_borders(t, top_c="CBD5E1", bot_c="CBD5E1", l_c="E2E8F0", r_c="E2E8F0", ih_c="E2E8F0"):
    tbl = t._tbl; sp = tbl.tblPr
    if sp is None:
        sp = parse_xml('<w:tblPr %s/>' % nsdecls("w")); tbl.insert(0, sp)
    for e in sp.findall(qn('w:tblBorders')): sp.remove(e)
    sp.append(parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="%s"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="%s"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '</w:tblBorders>' % (nsdecls("w"), top_c, bot_c, l_c, r_c, ih_c, ih_c)))

def no_borders(t):
    tbl = t._tbl; sp = tbl.tblPr
    if sp is None:
        sp = parse_xml('<w:tblPr %s/>' % nsdecls("w")); tbl.insert(0, sp)
    for e in sp.findall(qn('w:tblBorders')): sp.remove(e)

def set_cell_border(c, side="bottom", sz=8, color="1E3A5F"):
    tcPr = c._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = ET.SubElement(tcPr, qn('w:tcBorders'))
    else:
        for b in borders.findall(qn('w:' + side)):
            borders.remove(b)
    ET.SubElement(borders, qn('w:' + side)).attrib.update({
        qn('w:val'): 'single', qn('w:sz'): str(sz),
        qn('w:space'): '0', qn('w:color'): color
    })

def mk_table(doc, hdrs, rows, hdr_width_pct=None):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t)
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]
        mg(c, 80, 80, 150, 150)
        set_cell_border(c, "bottom", 12, "1E3A5F")
        cll(c, h, b=True, sz=9.5, clr="1E3A5F")
    for ri, rd in enumerate(rows):
        for ci, tx in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            mg(c, 60, 60, 150, 150)
            txt_clr = "1E293B" if ci == 0 else ("475569" if ri % 2 == 0 else "334155")
            cll(c, str(tx), b=(ci == 0), sz=8.5, clr=txt_clr)
    return t

def sec_title(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 0.3
    r = p.add_run("-"*60); r.font.size = Pt(1); r.font.color.rgb = RGBColor.from_string("93C5FD")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(8)
    r1 = p2.add_run(" %s " % num)
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = RGBColor.from_string("FFFFFF")
    r1._r.get_or_add_rPr().append(parse_xml('<w:shd %s w:fill="1D4ED8" w:val="clear"/>' % nsdecls("w")))
    rf(r1, ea="微软雅黑")
    r2 = p2.add_run("  %s" % title)
    r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = RGBColor.from_string("0F172A")
    rf(r2, ea="微软雅黑")

def info_box(doc, txt, sz=9, clr="1E3A5F"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    mg(c, 100, 100, 180, 180)
    tbl_borders(t, "60A5FA", "60A5FA", "60A5FA", "60A5FA", "60A5FA")
    tcPr = c._tc.get_or_add_tcPr()
    # Remove existing tcBorders before adding new one
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    borders = ET.SubElement(tcPr, qn('w:tcBorders'))
    ET.SubElement(borders, qn('w:left')).attrib.update({qn('w:val'): 'single', qn('w:sz'): '24', qn('w:space'): '0', qn('w:color'): '3B82F6'})
    cll(c, txt, sz=sz, clr=clr, lh=1.15)

def body_txt(doc, txt, sz=9, clr="334155", b=False, it=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(txt); r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor.from_string(clr); r.bold = b; r.italic = it; rf(r)

def dark_banner(doc, line1, line2):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    mg(c, 100, 100, 180, 180)
    tcPr = c._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    borders = ET.SubElement(tcPr, qn('w:tcBorders'))
    ET.SubElement(borders, qn('w:left')).attrib.update({qn('w:val'): 'single', qn('w:sz'): '36', qn('w:space'): '0', qn('w:color'): '1D4ED8'})
    ET.SubElement(borders, qn('w:bottom')).attrib.update({qn('w:val'): 'single', qn('w:sz'): '12', qn('w:space'): '0', qn('w:color'): '1D4ED8'})
    tbl_borders(t, "E2E8F0", "E2E8F0", "E2E8F0", "E2E8F0", "E2E8F0")
    cll(c, line1, b=True, sz=12, clr="1E3A5F", al=WD_ALIGN_PARAGRAPH.CENTER)
    cll(c, line2, sz=8, clr="3B82F6", al=WD_ALIGN_PARAGRAPH.CENTER)

# ---- WATERMARK ----
def add_wm(sec, img):
    hdr = sec.header; hdr.is_linked_to_previous = False
    for p in hdr.paragraphs: p.clear()
    p = hdr.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(img, width=Inches(3.5), height=Inches(2.2))
    drw = run._r.find(qn('w:drawing'))
    if drw is None: return
    il = drw.find(qn('wp:inline'))
    if il is None: return
    ext = il.find(qn('wp:extent'))
    cx = ext.get('cx', '3200400'); cy = ext.get('cy', '2011680')
    gr = il.find(qn('a:graphic'))
    px = str(max(0, 7938000//2 - int(cx)//2))
    py = str(max(0, 11226600//2 - int(cy)//2))
    ax = ET.fromstring((
        '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0"'
        ' behindDoc="1" locked="1" layoutInCell="0" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>%s</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>%s</wp:posOffset></wp:positionV>'
        '<wp:extent cx="%s" cy="%s"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="999" name="Watermark" descr="Logo"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '</wp:anchor>' % (px, py, cx, cy)
    ).encode('utf-8'))
    if gr is not None:
        ax.append(copy.deepcopy(gr))
        blp = ax.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if blp is not None:
            blp.append(ET.fromstring('<a:grayscl xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>'))
            al = ET.SubElement(blp, qn('a:alphaModFix'))
            al.set('amt', '18000')
    drw.remove(il); drw.append(ax)

add_wm(doc.sections[0], LOGO)

# ---- FOOTER ----
footer = doc.sections[0].footer
footer.is_linked_to_previous = False
pf = footer.paragraphs[0]; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_after = Pt(0)
r1 = pf.add_run("-"*60); r1.font.size = Pt(1); r1.font.color.rgb = RGBColor.from_string("CBD5E1")
pf2 = footer.add_paragraph()
pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf2.paragraph_format.space_before = Pt(2)
rl = pf2.add_run("BLIIOT Confidential  |  Page ")
rl.font.size = Pt(7.5); rl.font.color.rgb = RGBColor.from_string("94A3B8")
ru1 = pf2.add_run(); ru1._r.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w")))
ru2 = pf2.add_run(); ru2._r.append(parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls("w")))
ru3 = pf2.add_run(); ru3._r.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w")))

# ============ BUILD ============

# --- HEADER ---
ht = doc.add_table(rows=1, cols=2); ht.alignment = WD_TABLE_ALIGNMENT.CENTER; no_borders(ht)
tbl = ht._tbl; sp = tbl.tblPr
if sp is None: sp = parse_xml('<w:tblPr %s/>' % nsdecls("w")); tbl.insert(0, sp)
sp.append(parse_xml('<w:tblBorders %s><w:bottom w:val="single" w:sz="24" w:space="0" w:color="1E3A5F"/></w:tblBorders>' % nsdecls("w")))

tc = ht.rows[0].cells[0]; mg(tc, 0, 0, 0, 200)
p = tc.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
r = p.add_run("Zimbabwe ZODSAT\nTransformer Anti-Theft & Intrusion Prevention")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor.from_string("0F172A"); rf(r, ea="微软雅黑")
p2 = tc.add_paragraph(); p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.space_before = Pt(2)
r2 = p2.add_run("S275 LoRaWAN Custom Terminal Node — Product Specification")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor.from_string("1D4ED8"); rf(r2, ea="微软雅黑")
p3 = tc.add_paragraph()
r3 = p3.add_run("& Requirements Confirmation Document")
r3.bold = True; r3.font.size = Pt(11); r3.font.color.rgb = RGBColor.from_string("1D4ED8"); rf(r3, ea="微软雅黑")

lc = ht.rows[0].cells[1]; mg(lc, 0, 0, 0, 0)
pl = lc.paragraphs[0]; pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pl.add_run().add_picture(LOGO, width=Inches(1.5), height=Inches(0.94))

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# --- INFO TABLE ---
info = [
    ("Client", "SmartGrid Africa / ZODSAT (Zimbabwe)"),
    ("Trading Entity", "SmartGrid Africa — Transaction & Commercial Management"),
    ("Implementing Entity", "ZODSAT — Field Deployment & Operations"),
    ("Supplier", "Shenzhen Beilai Technology Co., Ltd. (BLIIOT)"),
    ("Purpose", "Final product specification for customer review — features, accessories, project timeline."),
    ("Version", "V5.0 — Updated Specification & Corrected Pricing"),
    ("Date", "27 May 2026"),
    ("Classification", "Confidential — For Authorized ZODSAT / SmartGrid Africa Personnel Only"),
]
it = doc.add_table(rows=len(info), cols=2); it.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(it)
for i, (l, v) in enumerate(info):
    cl = it.rows[i].cells[0]; mg(cl, 50, 50, 120, 120)
    cll(cl, l, b=True, sz=9, clr="1E293B")
    cv = it.rows[i].cells[1]; mg(cv, 50, 50, 120, 120)
    cll(cv, v, sz=9, clr="475569")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# --- INSTRUCTIONS ---
info_box(doc,
    "Customer Review Instructions\n\n"
    "This document presents the complete technical specifications, product features, design scope, "
    "integrated accessories, and project timeline for the Zimbabwe ZODSAT / SmartGrid Africa "
    "Transformer Anti-Theft and Intrusion Prevention Project, based on our technical evaluations and discussions.\n\n"
    "Please review every item carefully. If any information is missing, inaccurate, or inconsistent "
    "with your expectations, please mark it in this document or notify us in writing.\n\n"
    "If all information is correct, please sign the Customer Confirmation section at the end. "
    "The signed version will serve as the basis for prototype design, quotation, and project execution."
)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ====== SECTION 1 ======
sec_title(doc, "1", "Overall Technical Route")
mk_table(doc, ["Item", "Requirement / Current Understanding"], [
    ["Device Type", "LoRaWAN End Node / Terminal Node. BLIIOT does not require a LoRaWAN Gateway."],
    ["Communication", "LoRaWAN only (no 4G). Customer is building national LoRaWAN network for TAIS."],
    ["Existing Gateway", "Milesight Solar Gateway SG50-L08GL-868M PN 30W — already deployed."],
    ["Base Hardware", "BLIIOT S275 with LoRaWAN EU868 module. Large lithium battery, protective enclosure."],
    ["Integrated Scope", "S275 mainboard, LoRaWAN module, lithium battery, solar charging, relays, sensors, antenna, enclosure, customer logo."],
    ["Local Alarm Logic", "MUST support IF-THEN logic. Siren/strobe on sensor trigger even without LoRaWAN — NO cloud dependency."],
    ["Prototypes", "20 test prototypes. 2 engineering prototypes provided first."],
])

# ====== SECTION 2 ======
sec_title(doc, "2", "Key Hardware & I/O Requirements")
mk_table(doc, ["Module", "Specification"], [
    ["LoRaWAN Module", "EU868 (863–870 MHz); 8 multi-SF channels; ADR supported."],
    ["Data Path", "Node to Milesight SG50 Gateway to App Server to MQTT Cloud Platform."],
    ["Digital Input (DI)", "8 Dry Contact: door/hatch switches, PIR, vibration, power, tamper."],
    ["Analog Input (AI)", "2 channels, 4–20 mA or 0–10 V, 12-bit, for oil level / pressure."],
    ["Temperature", "Transformer oil/winding temp via PT100-to-4–20 mA transmitter."],
    ["Relay Output (DO)", "2 relays: DO1 = siren/strobe, DO2 = load shedding. 5 A @ 250 VAC / 30 VDC."],
    ["Local Logic Engine", "IF-THEN: PIR, door switch, vibration, power loss, temp, valve tamper."],
    ["External Antenna", "8 dBi fibre-glass antenna with 3 m cable."],        # FIXED: was 7dBi
    ["Configuration Tool", "PC software — equivalent/superior to Milesight UC300."],
    ["Branding", "Customer logo on final product enclosure."],
])

# ====== SECTION 3 ======
sec_title(doc, "3", "Sensor & Output Function List")
mk_table(doc, ["Function", "Description"], [
    ["Door / Access Hatch", "Magnetic switches for enclosure door / hatch monitoring."],
    ["PIR Motion Detection", "Human presence detection near transformer sites."],
    ["Vibration Sensor", "Accelerometer trigger for tampering or transformer movement."],
    ["Mains Power Status", "AC mains monitoring to detect power loss (theft indicator)."],
    ["Emergency / Tamper Loop", "Valve tampering detection for oil theft prevention."],
    ["PT100 Temperature", "Oil and winding temperature monitoring (via PT100-to-4–20mA converter)."],
    ["Oil Level / Pressure", "AI-based monitoring via 4–20 mA sensors."],
    ["Siren / Strobe (DO1)", "Local alarm — instant trigger on breach detection."],
    ["Load Shedding (DO2)", "Remote load shedding or system interlocking."],
    ["RS485 Relay Board", "Compatible with existing 8-channel RS485 relay board."],
    # NEW: IoT Lock integration
    ["IoT Cabinet Lock (NEW)", "Electronic cabinet lock with hidden backup key. Integrates with ThingsBoard for remote unlock, operator access logging, and maintenance audit trail."],
])

# ====== SECTION 4 ======
sec_title(doc, "4", "Power Supply, Battery, Solar & Enclosure")
mk_table(doc, ["Item", "Specification"], [
    ["Battery Backup", "Min. 72 hours during power outage (all sensors + siren active)."],
    # FIXED: Added 11.1V nominal voltage
    ["Battery Type", "11.1 V (nominal 3S Li-ion) / 12 V compatible. 5 Ah rechargeable with BMS."],
    ["Primary Power", "Turnkey solar: panel + Li-ion battery. Eliminates transformer surge damage."],
    ["Solar Input", "DC 9–36 V (S275 input range). BLIIOT provides panel specifications."],
    ["Installation Options", "(1) Outdoor/pole: panel mounted on TAIS unit with pole-mount stand. (2) Indoor/wall: panel on roof, cabled to unit."],
    ["Charging", "100 W MPPT charger. BMS, status detection, low-voltage alarm, overcharge protection."],
    ["Enclosure", "Professional protective enclosure, DIN-rail / wall-mount, compact."],
    ["Outdoor Protection", "Waterproof, dustproof, corrosion-resistant for all external connections."],
])

# ====== SECTION 5 ======
sec_title(doc, "5", "Platform System Requirements (ThingsBoard Boundary)")
body_txt(doc, "Note: Platform scope is separate from S275 hardware. Included for payload format, data fields, alarm rules, and integration interface reference.", sz=8.5, clr="64748B", it=True)
mk_table(doc, ["Item", "Requirement"], [
    ["Communication", "MQTT cloud with Milesight Gateway, LoRaWAN End Nodes, Application Server."],
    ["Alarm Notifications", "WhatsApp, SMS, Email. Future: phone call to police / response teams."],
    ["Map Dashboard", "Site map: green = online, red = offline. Alarm events show site location."],
    ["Front-end & Back-end", "Full UI + backend development to replace Akenza API."],
    ["IoT Lock Integration", "ThingsBoard-integrated: remote unlock, authorized operator log, maintenance audit trail."],
    ["Integration Boundary", "Scope, timeline, cost defined separately. HW provides payload definitions and data fields."],
])

# ====== SECTION 6 ======
sec_title(doc, "6", "Integrated Accessories & Commercial Scheme")
dark_banner(doc,
    "Development & Engineering NRE Fee: USD $10,000.00 (All-Inclusive)",
    "Covers: custom firmware, LoRaWAN integration, solar charging design, IoT lock wiring, 2 prototype samples, express air freight to Harare.")
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# FIXED: Updated 50W panel price from $19.5 to $9.2
# Recalculated: $80 + ($10.51 + $1.50 + $6.50 + $9.20 + $5.12) = $112.83
# Recalculated: $70 + ($10.51 + $1.50 + $6.50 + $9.20 + $5.12) = $102.83
# Recalculated: $65 + ($10.24 + $1.50 + $6.00 + $9.20 + $5.12) = $97.05

mk_table(doc, ["Qty", "Solar Panel", "Base Node (USD)", "Accessories* (USD)", "Total (EXW)"], [
    ["2,000", "30 W", "$80.00", "$36.63", "$116.63"],
    ["2,000", "40 W", "$80.00", "$39.13", "$119.13"],
    ["2,000", "50 W", "$80.00", "$32.83", "$112.83"],      # UPDATED
    ["5,000", "30 W", "$70.00", "$36.63", "$106.63"],
    ["5,000", "40 W", "$70.00", "$39.13", "$109.13"],
    ["5,000", "50 W", "$70.00", "$32.83", "$102.83"],      # UPDATED
    ["10,000", "30 W", "$65.00", "$35.85", "$100.85"],
    ["10,000", "40 W", "$65.00", "$38.35", "$103.35"],
    ["10,000", "50 W", "$65.00", "$32.05", "$97.05"],      # UPDATED
])

body_txt(doc, "*Accessories: 11.1 V / 5 Ah Li-ion Battery, 100 W MPPT Solar Charger, IoT Cabinet Lock (hidden key override), 8 dBi Antenna + 3 m Cable, Pole-Mount Stand.", sz=7.5, clr="64748B", it=True, sb=2, sa=1)
body_txt(doc, "Promotional (50 W): 2,000 @ $149.00  |  5,000 @ $144.00  |  10,000 @ $139.00", sz=8, clr="1D4ED8", b=True)

# NEW: Payment Terms
body_txt(doc, "", sz=4, sa=2)
p_pt = doc.add_paragraph()
p_pt.paragraph_format.space_before = Pt(2); p_pt.paragraph_format.space_after = Pt(2)
r_pt_title = p_pt.add_run("Payment Terms")
r_pt_title.bold = True; r_pt_title.font.size = Pt(10); r_pt_title.font.color.rgb = RGBColor.from_string("0F172A"); rf(r_pt_title, ea="微软雅黑")

pt_tbl = doc.add_table(rows=3, cols=2); pt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(pt_tbl)
for j, h in enumerate(["Term", "Condition"]):
    c = pt_tbl.rows[0].cells[j]; mg(c, 60, 60, 120, 120)
    set_cell_border(c, "bottom", 10, "1E3A5F")
    cll(c, h, b=True, sz=9, clr="1E3A5F")
pt_data = [
    ("Deposit", "50% of the total order value upon order confirmation."),
    ("Balance", "50% balance payable before shipment."),
]
for i, (l, v) in enumerate(pt_data):
    cl = pt_tbl.rows[i+1].cells[0]; mg(cl, 50, 50, 120, 120)
    cll(cl, l, b=True, sz=9, clr="1E293B")
    cv = pt_tbl.rows[i+1].cells[1]; mg(cv, 50, 50, 120, 120)
    cll(cv, v, sz=9, clr="475569")

body_txt(doc, "Quotation Term: EXW (Shenzhen, China). All prices in USD.", sz=8, clr="64748B", it=True, sb=2, sa=0)

# ====== SECTION 7 ======
sec_title(doc, "7", "Project Development & Delivery Schedule")
body_txt(doc, "Preliminary — subject to confirmation upon specification freeze and deposit.", sz=8, clr="B91C1C", it=True)
mk_table(doc, ["Phase", "Timeline"], [
    ["Schematic & PCB Modification", "4 weeks from specification freeze and deposit"],
    ["Prototype Production & Debugging (2 units)", "15 working days (incl. firmware porting)"],
    ["Prototype Delivery (Harare)", "International express air freight — TBC"],
    ["Mass Production (first 2,000 pcs)", "30 working days after prototype approval"],
    ["Monthly Capacity", "400–700 units / month (sustained, scalable)"],
])

# ====== SECTION 8 ======
sec_title(doc, "8", "Customer Confirmation & Signature")
info_box(doc,
    '"I / We confirm that we have reviewed this document. The requirements, technical specifications, '
    'and project understanding are complete and accurate as of the signing date."')
doc.add_paragraph().paragraph_format.space_after = Pt(6)

st = doc.add_table(rows=7, cols=2); st.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(st)
for j, h in enumerate(["Field", "Information to be Completed"]):
    c = st.rows[0].cells[j]; mg(c, 80, 80, 120, 120)
    set_cell_border(c, "bottom", 12, "1E3A5F")
    cll(c, h, b=True, sz=9, clr="1E3A5F")
sig_data = ["Customer Company", "Authorized Representative", "Position / Title", "Signature", "Date"]
for i, lbl in enumerate(sig_data):
    rn = i+1
    cl = st.rows[rn].cells[0]; cv = st.rows[rn].cells[1]
    mg(cl, 60, 60, 120, 120); cll(cl, lbl, b=True, sz=9, clr="1E293B")
    mg(cv, 60, 60, 120, 120); cll(cv, "_"*50, sz=7, clr="CBD5E1")

sr = st.rows[6]; sr.cells[0].merge(sr.cells[1])
mg(sr.cells[0], 60, 60, 120, 120)
set_cell_border(sr.cells[0], "top", 6, "CBD5E1")
cll(sr.cells[0], "Company Stamp / Seal (if applicable)", sz=8.5, clr="64748B")

# --- END ---
pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
pe.paragraph_format.space_before = Pt(20)
re = pe.add_run("--- End of Document ---")
re.font.size = Pt(10); re.font.color.rgb = RGBColor.from_string("94A3B8"); re.italic = True; rf(re)

doc.save(OUT)
print("OK: " + OUT)
