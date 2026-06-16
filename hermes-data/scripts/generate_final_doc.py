#!/c/Users/Admin/AppData/Local/Programs/Python/Python313/python.exe
# -*- coding: utf-8 -*-
"""Generate the final updated document for Zimbabwe ZODSAT project"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, copy

LOGO_PATH = r"C:\Users\Admin\Desktop\Working\津巴布尔客户\签署文件\钡铼英文LOGO.png"

doc = Document()

# ===== Page Setup =====
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== Helper Functions =====

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set cell margins"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    # Remove existing 
    existing = tcPr.findall(qn('w:tcMar'))
    for e in existing:
        tcPr.remove(e)
    tcPr.append(tcMar)

def add_formatted_para(cell, text, bold=False, font_size=10, color="1E293B", font_name="Segoe UI", alignment=None, space_after=0):
    """Add a formatted paragraph to a cell"""
    p = cell.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font_name
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_section_heading(doc, number, title):
    """Add a styled section heading"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    # Add a colored bar background effect via shading on the paragraph
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("1D4ED8")
    run.font.name = "Segoe UI"
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_info_row(table, label, value, label_bg="F8FAFC"):
    """Add an info row to the header table"""
    row = table.add_row()
    # Label cell
    cell_label = row.cells[0]
    set_cell_shading(cell_label, label_bg)
    set_cell_margins(cell_label, 80, 80, 150, 150)
    add_formatted_para(cell_label, label, bold=True, font_size=10, color="1E293B")
    # Value cell
    cell_val = row.cells[1]
    set_cell_margins(cell_val, 80, 80, 150, 150)
    add_formatted_para(cell_val, value, bold=False, font_size=10, color="1E293B")
    if label_bg == "F8FAFC":
        set_cell_shading(cell_val, "FFFFFF")
    return row

def set_table_borders(table, top_color="CBD5E1", bottom_color="CBD5E1", inside_color="E2E8F0"):
    """Set table borders"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{top_color}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{bottom_color}"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{inside_color}"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders element
    existing = tblPr.findall(qn('w:tblBorders'))
    for e in existing:
        tblPr.remove(e)
    tblPr.append(borders)

def set_table_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_blue_white_table(doc, headers, rows, col_widths=None):
    """Add a nice styled table with blue header and white rows"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "EFF6FF")
        set_cell_margins(cell, 100, 100, 150, 150)
        add_formatted_para(cell, header, bold=True, font_size=9, color="1E293B")
    
    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            bg = "F8FAFC" if r % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)
            set_cell_margins(cell, 80, 80, 150, 150)
            add_formatted_para(cell, str(cell_text), bold=False, font_size=9, color="334155")
    
    return table

# ===== WATERMARK via Header =====
def add_watermark_to_section(section, image_path):
    """Add watermark image to header of a section"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    
    # Add image as watermark (centered, semi-transparent effect)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(2.8), height=Inches(1.76))
    
    # Position the picture - we need to make it a watermark-like placement
    # Get the inline picture element
    drawing = run._r.findall(qn('w:drawing'))[0] if run._r.findall(qn('w:drawing')) else None
    
    if drawing:
        # We'll set the image as positioned behind text
        # For simplicity, set it as a watermark in the center
        anchor_xml = f'''<w:anchor xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main"
            xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
            xmlns:mv="urn:schemas-microsoft-com:mac:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
            xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
            xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
            xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
            xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
            mc:Ignorable="w14 wp14"
            distT="0" distB="0" distL="0" distR="0"
            simplePos="0" relativeHeight="0" behindDoc="1"
            locked="0" layoutInCell="1" allowOverlap="1"
            wp14:anchorId="00000000">
        </w:anchor>'''
        # We'll just use the inline image for simplicity - it's in the header which works as a watermark

# Add watermark to first section
add_watermark_to_section(doc.sections[0], LOGO_PATH)

# ===== TITLE HEADER TABLE =====
title_table = doc.add_table(rows=1, cols=2)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Remove borders from title table
tbl = title_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:bottom w:val="single" w:sz="18" w:space="0" w:color="0F172A"/>'
    f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
existing = tblPr.findall(qn('w:tblBorders'))
for e in existing:
    tblPr.remove(e)
tblPr.append(borders)

# Title cell
title_cell = title_table.rows[0].cells[0]
title_cell.width = Inches(4.5)
p1 = title_cell.paragraphs[0]
p1.paragraph_format.space_after = Pt(4)
run_title = p1.add_run("Zimbabwe ZODSAT Transformer Anti-Theft\nand Intrusion Prevention")
run_title.bold = True
run_title.font.size = Pt(22)
run_title.font.color.rgb = RGBColor.from_string("0F172A")
run_title.font.name = "Segoe UI"

p2 = title_cell.paragraphs[0] if False else title_cell.add_paragraph()
# Actually create a new paragraph
p2 = title_cell.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
p2.paragraph_format.space_before = Pt(0)
run_sub = p2.add_run("S275 LoRaWAN Custom Terminal Node – Product Specification &\nCustomer Requirements Confirmation Document")
run_sub.bold = True
run_sub.font.size = Pt(15)
run_sub.font.color.rgb = RGBColor.from_string("1D4ED8")
run_sub.font.name = "Segoe UI"

# Logo cell
logo_cell = title_table.rows[0].cells[1]
logo_cell.width = Inches(2.5)
pl = logo_cell.paragraphs[0]
pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pl.paragraph_format.space_after = Pt(0)
pl.paragraph_format.space_before = Pt(0)
run_logo = pl.add_run()
run_logo.add_picture(LOGO_PATH, width=Inches(1.5), height=Inches(0.94))

# Space after title
doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ===== INFO TABLE =====
info_table = doc.add_table(rows=1, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(info_table)

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Inches(2.2)
    row.cells[1].width = Inches(4.8)

# Remove the default empty row and add proper rows
info_table._tbl.findall(qn('w:tr'))  # get rows
# Remove first (default) row
first_row = info_table.rows[0]
first_row._tr.getparent().remove(first_row._tr)

info_rows = [
    ("Prepared for", "ZODSAT (Zimbabwe)"),
    ("Prepared by", "Shenzhen Beilai Technology Co., Ltd. (BLIIOT)"),
    ("Document Purpose", "Technical specification for the final product – detailing all product features, functions, design scope, integrated accessories, and project timeline for customer review and confirmation."),
    ("Document Status", "Updated Specification Draft v2.0"),
    ("Prepared Date", "27 May 2026"),
]
for label, value in info_rows:
    add_info_row(info_table, label, value, 
                 label_bg="F8FAFC" if info_rows.index((label, value)) % 2 == 0 else "FFFFFF")

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ===== CUSTOMER REVIEW INSTRUCTIONS =====
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Customer Review Instructions")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor.from_string("1E293B")
run.font.name = "Segoe UI"

instruction_bg_table = doc.add_table(rows=1, cols=1)
instruction_bg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_shading(instruction_bg_table.rows[0].cells[0], "EFF6FF")
set_cell_margins(instruction_bg_table.rows[0].cells[0], 140, 140, 180, 180)
cell = instruction_bg_table.rows[0].cells[0]
add_formatted_para(cell, 
    "This document is prepared for your review and confirmation. It summarizes the complete technical specifications, "
    "product features, design scope, integrated accessories, and project timeline for the Zimbabwe ZODSAT Transformer "
    "Anti-Theft and Intrusion Prevention Project based on our discussions and technical evaluations.\n\n"
    "Please review every item carefully and confirm whether this document fully and accurately covers all your requirements. "
    "If any requirement is missing, inaccurate, unclear, or inconsistent with your expectations, please mark it in this "
    "document or inform us in writing. We will update the document promptly.\n\n"
    "If all information is correct and complete, please complete and sign the Customer Confirmation section at the end of "
    "this document. The signed version will be used as the basis for subsequent technical evaluation, prototype design, "
    "quotation, and project communication. This document is a requirements confirmation document and is not a final "
    "commercial contract unless otherwise agreed in writing.",
    bold=False, font_size=9, color="1E293B"
)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 1: Overall Technical Route =====
add_section_heading(doc, "1", "Overall Technical Route for Confirmation")

tech_headers = ["Item", "Current Understanding / Requirement for Your Confirmation"]
tech_rows = [
    ["Device Type", "LoRaWAN End Node / Terminal Node. BLIIOT does not need to provide a LoRaWAN Gateway as confirmed by the customer."],
    ["Communication Method", "LoRaWAN only, without 4G. Local 4G network is unstable; previous system experienced alarm delays due to 4G latency. Customer is building a national LoRaWAN network for TAIS."],
    ["Existing Gateway", "Milesight Solar Gateway SG50-L08GL-868M PN 30W (already deployed by customer)."],
    ["Base Hardware Route", "Use BLIIOT S275 as the base device, replace the original 4G module with a LoRaWAN EU868 module, integrate a large-capacity lithium battery into a professional protective enclosure."],
    ["Integrated Unit Scope", "S275 main board, LoRaWAN module, lithium battery, solar charging management, relay outputs, sensor terminals, external antenna interface, protective enclosure, and customer logo branding."],
    ["Local Alarm Requirement", "Terminal MUST support local IF-THEN logic. When network/LoRaWAN is abnormal, siren/strobe must activate immediately upon sensor trigger – no cloud dependency."],
    ["Prototype Quantity", "20 test prototypes confirmed. BLIIOT will first provide 2 engineering prototypes for evaluation."],
]
add_blue_white_table(doc, tech_headers, tech_rows)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 2: Key Hardware and I/O =====
add_section_heading(doc, "2", "Key Hardware and I/O Requirements for Confirmation")

io_headers = ["Module", "Requirement for Your Confirmation"]
io_rows = [
    ["LoRaWAN Communication", "EU868 / EU863–870 for Zimbabwe; frequency 863–870 MHz; 8 multi-SF channels (868.1–867.9 MHz); ADR supported to optimize battery life and range."],
    ["LoRaWAN Data Path", "Node reports data via customer's Milesight SG50 Gateway → Application Server → MQTT Cloud Platform. Integration boundaries confirmed during joint debugging."],
    ["Digital Input (DI)", "8 Dry Contact inputs for: Door/Access Hatch magnetic switches, PIR/Vibration, Mains power status, Emergency/Tamper loop, etc."],
    ["Analog Input (AI)", "2 AI channels, 4–20 mA or 0–10 V, 12-bit resolution, for oil level / pressure monitoring."],
    ["PT100 Temperature", "Transformer oil/winding temperature monitoring. S275 does not support direct PT100; uses PT100-to-4–20mA transmitter solution."],
    ["DO / Relay", "2 relay outputs: DO1 for local siren/strobe, DO2 for optional remote load shedding/interlocking. Rating: 5A@250VAC or 5A@30VDC."],
    ["Local Logic / Alarm", "IF-THEN engine: trigger sources include PIR, door magnetic switch, vibration, Power Loss, temperature threshold, Valve Tampering. Cloud rules also used."],
    ["External Antenna", "7dBi antenna with 3m cable for installation flexibility inside transformer rooms."],
    ["Configuration Software", "PC-based configuration tool, equivalent to or better than Milesight UC300 capabilities."],
    ["Branding / Appearance", "Final product will carry the customer's logo on the enclosure."],
]
add_blue_white_table(doc, io_headers, io_rows)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 3: Sensor and Output Function List =====
add_section_heading(doc, "3", "Sensor and Output Function List for Confirmation")

sensor_headers = ["Function Category", "Content for Your Confirmation"]
sensor_rows = [
    ["Door / Access Hatch Sensor", "Magnetic switches for enclosure door/access hatch monitoring."],
    ["PIR / Motion Detection", "PIR motion sensors detecting human presence near transformers."],
    ["Vibration Sensor", "Accelerometer-based trigger detecting physical tampering or transformer movement."],
    ["Mains / Transformer Power Status", "Mains power monitoring to detect power loss (potential theft indicator)."],
    ["Emergency / Tamper Loop", "Emergency loop and valve tampering detection for oil theft prevention."],
    ["PT100 Temperature", "Oil and winding temperature monitoring via PT100-to-4–20mA transmitter."],
    ["Oil Level / Pressure", "AI-based oil level and pressure monitoring via 4–20mA sensors."],
    ["Siren / Strobe Alarm", "Local alarm output (DO1) – triggers siren/strobe instantly upon breach detection."],
    ["Load Shedding / Interlocking", "Secondary output (DO2) for remote load shedding or system interlocking."],
    ["RS485 Relay Board", "Compatible with customer's existing 8-channel RS485 relay board if needed."],
]
add_blue_white_table(doc, sensor_headers, sensor_rows)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 4: Power Supply, Battery, Solar =====
add_section_heading(doc, "4", "Power Supply, Battery, Solar Power, and Enclosure Requirements")

power_headers = ["Item", "Technical Requirement / Current Understanding for Confirmation"]
power_rows = [
    ["Battery Capacity & Backup", "Minimum 72 hours operation during power outage including all security sensors, temperature sensor, vibration sensor, and siren. Customer's existing 5Ah Li-ion battery provides ~4 days backup without solar."],
    ["Battery Type", "Lithium battery pack (12V 5Ah rechargeable). BMS included for overcharge/over-discharge protection."],
    ["Solar Power as Main Supply", "Turnkey solar power solution. Solar panel + lithium battery replaces previous transformer-tap power (eliminates surge damage, SPD/OVP costs)."],
    ["Solar Voltage", "Solar output must meet S275 input range: DC 9–36V. BLIIOT to provide panel specs for customer procurement."],
    ["Solar Installation Forms", "Option 1: Outdoor/Pole Mount – panel fixed on TAIS unit. Option 2: Indoor/Wall Mount – panel on roof/external structure, cabled to unit."],
    ["Solar Charging Management", "MPPT intelligent solar charger (100W) included. Features: BMS, charging status detection, low-voltage alarm, overcharge/discharge protection, solar input reporting."],
    ["Enclosure / Transponder Box", "Professional protective enclosure, DIN-rail or wall mount compatible. Size optimized – as compact as possible while housing main board, battery, and relay outputs."],
    ["Outdoor Cabling & Protection", "Waterproof, dustproof, corrosion-resistant design for all external connections: antenna cable, sensor cables, siren output, and solar panel cable."],
]
add_blue_white_table(doc, power_headers, power_rows)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 5: Platform System Requirements =====
add_section_heading(doc, "5", "Platform System Requirements and Hardware Project Boundary (ThingsBoard)")

p_note = doc.add_paragraph()
p_note.paragraph_format.space_after = Pt(6)
run_note = p_note.add_run(
    "In addition to the LoRaWAN terminal node, the customer has proposed a complete platform development requirement. "
    "This is not a function of the S275 hardware itself, but it affects payload format, device data fields, alarm rules, "
    "and integration interfaces. This section documents the platform scope for reference."
)
run_note.font.size = Pt(9)
run_note.font.color.rgb = RGBColor.from_string("64748B")
run_note.font.name = "Segoe UI"

plat_headers = ["Platform Item", "Customer Platform Requirement / Technical Boundary"]
plat_rows = [
    ["Platform Communication Path", "MQTT cloud platform communicating with Milesight Gateway, LoRaWAN End Nodes, and Application Server."],
    ["Alarm Notification", "Alarms via WhatsApp, SMS, Email. Future extension to phone call alerts to police/rapid response teams."],
    ["Map Dashboard", "Control center map view: green = online, red = offline. Alarm events highlighted with specific site location."],
    ["Front-end & Back-end Development", "Complete UI and backend development to replace current Akenza API approach."],
    ["Platform Boundary", "Platform scope, interfaces, development timeline, cost, and responsibility to be defined separately. Hardware project provides: clear payload definitions, device data fields, and integration support."],
]
add_blue_white_table(doc, plat_headers, plat_rows)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 6: Integrated Accessories & Commercial Scheme =====
add_section_heading(doc, "6", "Integrated Accessories Breakdown & Commercial Scheme")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Development & Engineering NRE Fee: USD $10,000.00 (All-Inclusive)")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string("1D4ED8")
run.font.name = "Segoe UI"

p_note_nre = doc.add_paragraph()
p_note_nre.paragraph_format.space_after = Pt(8)
run_nre = p_note_nre.add_run(
    "NRE covers: custom software modifications, LoRaWAN module integration, solar charging management design, "
    "IoT lock controller wiring, 2 fully functional prototype samples, and express air freight to Harare, Zimbabwe."
)
run_nre.font.size = Pt(9)
run_nre.font.color.rgb = RGBColor.from_string("64748B")
run_nre.font.name = "Segoe UI"

# Pricing table
price_headers = ["Quantity", "Solar Panel Spec", "Base Custom Node (USD)", "Accessories* Subtotal", "Final Package Price (EXW)"]
price_rows = [
    ["2,000 PCS", "30W Panel", "$80.00", "$36.63", "$116.63"],
    ["2,000 PCS", "40W Panel", "$80.00", "$39.13", "$119.13"],
    ["2,000 PCS", "50W Panel (Standard)", "$80.00", "$43.13", "$123.13 (Promo: $149.00)"],
    ["5,000 PCS", "30W Panel", "$70.00", "$36.63", "$106.63"],
    ["5,000 PCS", "40W Panel", "$70.00", "$39.13", "$109.13"],
    ["5,000 PCS", "50W Panel (Standard)", "$70.00", "$43.13", "$113.13 (Promo: $144.00)"],
    ["10,000 PCS", "30W Panel", "$65.00", "$35.85", "$100.85"],
    ["10,000 PCS", "40W Panel", "$65.00", "$38.35", "$103.35"],
    ["10,000 PCS", "50W Panel (Standard)", "$65.00", "$42.35", "$107.35 (Promo: $139.00)"],
]
add_blue_white_table(doc, price_headers, price_rows)

# Accessories note
p_acc = doc.add_paragraph()
p_acc.paragraph_format.space_before = Pt(6)
p_acc.paragraph_format.space_after = Pt(4)
run_acc = p_acc.add_run(
    "*Accessories Subtotal includes specialized components: 12V 5Ah Lithium Battery Pack, 100W MPPT Intelligent Solar Charger, "
    "12V DC IoT Cabinet Electronic Lock with hidden backup physical key override, 8dBi Fiberglass Antenna + 3m Cable, "
    "and Heavy-duty Pole-Mount Stand."
)
run_acc.italic = True
run_acc.font.size = Pt(8)
run_acc.font.color.rgb = RGBColor.from_string("64748B")
run_acc.font.name = "Segoe UI"

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ===== SECTION 7: Project Schedule =====
add_section_heading(doc, "7", "Project Development, Design, and Delivery Schedule")

p_sched_note = doc.add_paragraph()
p_sched_note.paragraph_format.space_after = Pt(6)
run_note = p_sched_note.add_run(
    "The following timeline is provided for your reference and is subject to confirmation upon specification freeze and deposit."
)
run_note.font.size = Pt(9)
run_note.font.color.rgb = RGBColor.from_string("64748B")
run_note.font.name = "Segoe UI"

sched_headers = ["Project Development Phase", "Timeline / Target Capacity"]
sched_rows = [
    ["Schematic & PCB Hardware Modification", "4 Weeks from Specification Freeze and Deposit Confirmation"],
    ["Prototype Production & Debugging (2 Units)", "15 Working Days (includes firmware porting)"],
    ["Prototype Site Deployment Test (Harare)", "Shipped via international express air freight for customer evaluation"],
    ["Mass Production Lead Time (First Batch of 2,000 PCS)", "30 Working Days following official prototype approval"],
    ["Monthly Continuous Output Capacity Rate", "400 to 700 units per month (sustained capability)"],
]
add_blue_white_table(doc, sched_headers, sched_rows)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ===== SECTION 8: Customer Confirmation and Signature =====
add_section_heading(doc, "8", "Customer Confirmation and Signature")

p_confirm = doc.add_paragraph()
p_confirm.paragraph_format.space_after = Pt(10)
run_confirm = p_confirm.add_run(
    '"I / We confirm that we have reviewed this document. The requirements, technical specifications, features, '
    'and project understanding listed in this document are complete and accurate as of the signing date. '
    'This signed confirmation may be used as the basis for subsequent technical evaluation, prototype design, '
    'quotation, and project communication."'
)
run_confirm.italic = True
run_confirm.font.size = Pt(10)
run_confirm.font.color.rgb = RGBColor.from_string("1E293B")
run_confirm.font.name = "Segoe UI"

# Signature table
sig_table = doc.add_table(rows=6, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(sig_table, top_color="CBD5E1", bottom_color="CBD5E1", inside_color="E2E8F0")

sig_fields = [
    ("Field", "Information to be Completed by Customer"),
    ("Customer Company", ""),
    ("Authorized Representative Name", ""),
    ("Position / Title", ""),
    ("Signature", ""),
    ("Date", ""),
]

# Set first row as header
for i, (label, value) in enumerate(sig_fields):
    if i == 0:
        # Header row
        for j, (txt) in enumerate([label, value]):
            cell = sig_table.rows[i].cells[j]
            set_cell_shading(cell, "EFF6FF")
            set_cell_margins(cell, 100, 100, 150, 150)
            add_formatted_para(cell, txt, bold=True, font_size=10, color="1E293B")
    else:
        cell_l = sig_table.rows[i].cells[0]
        set_cell_shading(cell_l, "F8FAFC")
        set_cell_margins(cell_l, 100, 100, 150, 150)
        add_formatted_para(cell_l, label, bold=True, font_size=10, color="1E293B")
        
        cell_v = sig_table.rows[i].cells[1]
        set_cell_margins(cell_v, 100, 100, 150, 150)
        if value:
            add_formatted_para(cell_v, value, bold=False, font_size=10, color="1E293B")
        else:
            add_formatted_para(cell_v, "__", bold=False, font_size=14, color="CBD5E1")

# Add stamp/seal row
stamp_row = sig_table.add_row()
stamp_row.cells[0].merge(stamp_row.cells[1])
set_cell_margins(stamp_row.cells[0], 100, 100, 150, 150)
set_cell_shading(stamp_row.cells[0], "FAFAFA")
add_formatted_para(stamp_row.cells[0], "Company Stamp / Seal (if applicable)", bold=False, font_size=10, color="64748B")

# ===== FOOTNOTE =====
doc.add_paragraph().paragraph_format.space_after = Pt(6)
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run("— End of Document —")
run_footer.font.size = Pt(10)
run_footer.font.color.rgb = RGBColor.from_string("94A3B8")
run_footer.font.name = "Segoe UI"

# ===== Add watermark to all sections =====
# Since python-docx creates one section by default, the existing watermark is enough
# But let's make sure the header shows on first page
section = doc.sections[0]
section.different_first_page_header_footer = False

# ===== SAVE =====
output_path = r"C:\Users\Admin\Desktop\Working\津巴布尔客户\签署文件\更新后新文件.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print("Done!")
