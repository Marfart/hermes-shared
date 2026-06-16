#!/c/Users/Admin/AppData/Local/Programs/Python/Python313/python.exe
# -*- coding: utf-8 -*-
"""v5.1-CN-FINAL - Polished Chinese document matching reference style"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree as ET
import copy

LOGO = r"C:\Users\Admin\Desktop\Working\津巴布尔客户\签署文件\钡铼英文LOGO.png"
OUT = r"C:\Users\Admin\AppData\Local\hermes\更新后新文件_中文版.docx"

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(1.65); sec.bottom_margin = Cm(1.65)
    sec.left_margin = Cm(1.6); sec.right_margin = Cm(1.6)

# ---- HELPERS ----

def mg(c, t=60, b=60, l=120, r=120):
    tcPr = c._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcMar')): tcPr.remove(e)
    tcPr.append(parse_xml('<w:tcMar %s><w:top w:w="%d" w:type="dxa"/><w:bottom w:w="%d" w:type="dxa"/><w:left w:w="%d" w:type="dxa"/><w:right w:w="%d" w:type="dxa"/></w:tcMar>' % (nsdecls("w"), t, b, l, r)))

def rf(run, nm="Calibri", ea="微软雅黑"):
    rPr = run._r.get_or_add_rPr()
    for e in rPr.findall(qn('w:rFonts')): rPr.remove(e)
    rPr.append(parse_xml('<w:rFonts %s w:ascii="%s" w:hAnsi="%s" w:eastAsia="%s"/>' % (nsdecls("w"), nm, nm, ea)))

def cll(c, txt, b=False, sz=10, clr="1E293B", al=None, lh=1.15):
    p = c.paragraphs[0]; p.clear()
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = lh
    if al: p.alignment = al
    r = p.add_run(txt); r.bold = b; r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor.from_string(clr); rf(r)

def tbl_borders(t, top_c="CBD5E1", bot_c="CBD5E1", l_c="E2E8F0", r_c="E2E8F0", ih_c="E2E8F0"):
    tbl = t._tbl; sp = tbl.tblPr
    if sp is None:
        sp = parse_xml('<w:tblPr %s/>' % nsdecls("w")); tbl.insert(0, sp)
    for e in sp.findall(qn('w:tblBorders')): sp.remove(e)
    sp.append(parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="%s"/>'
        '</w:tblBorders>' % (nsdecls("w"), top_c, bot_c, l_c, r_c, ih_c, ih_c)))

def no_borders(t):
    tbl = t._tbl; sp = tbl.tblPr
    if sp is None:
        sp = parse_xml('<w:tblPr %s/>' % nsdecls("w")); tbl.insert(0, sp)
    for e in sp.findall(qn('w:tblBorders')): sp.remove(e)

def set_cell_border(c, side="bottom", sz=8, color="1F4E79"):
    tcPr = c._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = ET.SubElement(tcPr, qn('w:tcBorders'))
    else:
        for b in borders.findall(qn('w:' + side)):
            borders.remove(b)
    ET.SubElement(borders, qn('w:' + side)).attrib.update({
        qn('w:val'): 'single', qn('w:sz'): str(sz),
        qn('w:space'): '0', qn('w:color'): color
    })

def mk_table(doc, hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t)
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]
        mg(c, 60, 60, 120, 120)
        set_cell_border(c, "bottom", 10, "1F4E79")
        cll(c, h, b=True, sz=9, clr="1F4E79")
    for ri, rd in enumerate(rows):
        for ci, tx in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            mg(c, 50, 50, 120, 120)
            txt_clr = "1E293B" if ci == 0 else ("475569" if ri % 2 == 0 else "334155")
            cll(c, str(tx), b=(ci == 0), sz=8.5, clr=txt_clr)
    return t

def sec_title(doc, num, title):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(18); pf.space_after = Pt(6)
    pf.line_spacing = 1.2
    r = p.add_run("%s.  %s" % (num, title))
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string("1F4E79")
    rf(r)

def info_box(doc, txt, sz=9, clr="1F4E79"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    mg(c, 80, 80, 150, 150)
    tbl_borders(t, "CBD5E1", "CBD5E1", "CBD5E1", "CBD5E1", "CBD5E1")
    tcPr = c._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    borders = ET.SubElement(tcPr, qn('w:tcBorders'))
    ET.SubElement(borders, qn('w:left')).attrib.update({qn('w:val'): 'single', qn('w:sz'): '20', qn('w:space'): '0', qn('w:color'): '1F4E79'})
    cll(c, txt, sz=sz, clr=clr, lh=1.15)

def body_txt(doc, txt, sz=9, clr="334155", b=False, it=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(txt); r.font.size = Pt(sz)
    r.font.color.rgb = RGBColor.from_string(clr); r.bold = b; r.italic = it; rf(r)

def dark_banner(doc, line1, line2):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    mg(c, 80, 80, 150, 150)
    tcPr = c._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    borders = ET.SubElement(tcPr, qn('w:tcBorders'))
    ET.SubElement(borders, qn('w:left')).attrib.update({qn('w:val'): 'single', qn('w:sz'): '30', qn('w:space'): '0', qn('w:color'): '1F4E79'})
    ET.SubElement(borders, qn('w:bottom')).attrib.update({qn('w:val'): 'single', qn('w:sz'): '10', qn('w:space'): '0', qn('w:color'): '1F4E79'})
    tbl_borders(t, "E2E8F0", "E2E8F0", "E2E8F0", "E2E8F0", "E2E8F0")
    cll(c, line1, b=True, sz=11, clr="1F4E79", al=WD_ALIGN_PARAGRAPH.CENTER)
    cll(c, line2, sz=7.5, clr="3B82F6", al=WD_ALIGN_PARAGRAPH.CENTER)

# ---- WATERMARK ----
def add_wm(sec, img):
    hdr = sec.header; hdr.is_linked_to_previous = False
    for p in hdr.paragraphs: p.clear()
    p = hdr.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(img, width=Inches(3.5), height=Inches(2.2))
    drw = run._r.find(qn('w:drawing'))
    if drw is None: return
    il = drw.find(qn('wp:inline'))
    if il is None: return
    ext = il.find(qn('wp:extent'))
    cx = ext.get('cx', '3200400'); cy = ext.get('cy', '2011680')
    gr = il.find(qn('a:graphic'))
    px = str(max(0, 7938000//2 - int(cx)//2))
    py = str(max(0, 11226600//2 - int(cy)//2))
    ax = ET.fromstring((
        '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0"'
        ' behindDoc="1" locked="1" layoutInCell="0" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>%s</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>%s</wp:posOffset></wp:positionV>'
        '<wp:extent cx="%s" cy="%s"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="999" name="Watermark" descr="Logo"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '</wp:anchor>' % (px, py, cx, cy)
    ).encode('utf-8'))
    if gr is not None:
        ax.append(copy.deepcopy(gr))
        blp = ax.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        if blp is not None:
            blp.append(ET.fromstring('<a:grayscl xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>'))
            al = ET.SubElement(blp, qn('a:alphaModFix'))
            al.set('amt', '18000')
    drw.remove(il); drw.append(ax)

add_wm(doc.sections[0], LOGO)

# ---- FOOTER ----
footer = doc.sections[0].footer
footer.is_linked_to_previous = False
pf = footer.paragraphs[0]; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.paragraph_format.space_after = Pt(0)
pf.paragraph_format.space_before = Pt(0)
r1 = pf.add_run("—"*60); r1.font.size = Pt(1); r1.font.color.rgb = RGBColor.from_string("CBD5E1")
pf2 = footer.add_paragraph()
pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf2.paragraph_format.space_before = Pt(2)
rl = pf2.add_run("钡铼技术 机密文件  |  第 ")
rl.font.size = Pt(7); rl.font.color.rgb = RGBColor.from_string("94A3B8")
ru1 = pf2.add_run(); ru1._r.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w")))
ru2 = pf2.add_run(); ru2._r.append(parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls("w")))
ru3 = pf2.add_run(); ru3._r.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w")))
rl2 = pf2.add_run(" 页")
rl2.font.size = Pt(7); rl2.font.color.rgb = RGBColor.from_string("94A3B8")

# ============ BUILD CHINESE ============

# --- HEADER ---
ht = doc.add_table(rows=1, cols=2); ht.alignment = WD_TABLE_ALIGNMENT.CENTER; no_borders(ht)
tbl = ht._tbl; sp = tbl.tblPr
if sp is None: sp = parse_xml('<w:tblPr %s/>' % nsdecls("w")); tbl.insert(0, sp)
sp.append(parse_xml('<w:tblBorders %s><w:bottom w:val="single" w:sz="20" w:space="0" w:color="1F4E79"/></w:tblBorders>' % nsdecls("w")))

tc = ht.rows[0].cells[0]; mg(tc, 0, 0, 0, 150)
p = tc.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
r = p.add_run("津巴布韦 ZODSAT\n变压器防盗防入侵项目")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string("0F172A"); rf(r, ea="微软雅黑")
p2 = tc.add_paragraph(); p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.space_before = Pt(1)
r2 = p2.add_run("S275 LoRaWAN 定制终端节点 — 产品规格书")
r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor.from_string("1F4E79"); rf(r2, ea="微软雅黑")
p3 = tc.add_paragraph()
r3 = p3.add_run("暨需求确认文件")
r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = RGBColor.from_string("1F4E79"); rf(r3, ea="微软雅黑")

lc = ht.rows[0].cells[1]; mg(lc, 0, 0, 0, 0)
pl = lc.paragraphs[0]; pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pl.add_run().add_picture(LOGO, width=Inches(1.3), height=Inches(0.82))

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# --- INFO TABLE ---
info = [
    ("客户名称", "SmartGrid Africa / ZODSAT（津巴布韦）"),
    ("交易主体", "SmartGrid Africa — 负责交易与商务管理"),
    ("实施主体", "ZODSAT — 负责现场部署与运维"),
    ("供应商", "深圳钡铼技术有限公司（BLIIOT）"),
    ("文档目的", "最终产品规格确认 — 功能、配件、项目时间表"),
    ("版本", "V5.1-CN — 最终版"),
    ("日期", "2026年5月27日"),
    ("密级", "机密 — 仅限 ZODSAT / SmartGrid Africa 授权人员查阅"),
]
it = doc.add_table(rows=len(info), cols=2); it.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(it)
for i, (l, v) in enumerate(info):
    cl = it.rows[i].cells[0]; mg(cl, 40, 40, 100, 100)
    cll(cl, l, b=True, sz=8.5, clr="1E293B")
    cv = it.rows[i].cells[1]; mg(cv, 40, 40, 100, 100)
    cll(cv, v, sz=8.5, clr="475569")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# --- INSTRUCTIONS ---
info_box(doc,
    "客户审核说明\n\n"
    "本文件详细列明了津巴布韦 ZODSAT / SmartGrid Africa 变压器防盗防入侵项目的完整技术规格、产品功能、设计范围、"
    "集成配件及项目时间表，内容基于双方技术评估与沟通讨论而制定。\n\n"
    "请逐项仔细审核。如有任何遗漏、错误或与您预期不一致之处，请在本文件中标注或书面通知我方，我方将及时更新。\n\n"
    '如所有信息确认无误，请在本文件末尾的"客户确认与签字"处签字。签字版本将作为样机设计、报价及项目执行的基础依据。'
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ====== SECTION 1 ======
sec_title(doc, "一", "总体技术路线")
mk_table(doc, ["项目", "当前理解 / 客户确认要求"], [
    ["设备类型", "LoRaWAN 终端节点。我方无需提供 LoRaWAN 网关。"],
    ["通信方式", "仅使用 LoRaWAN，不使用 4G。客户正在建设全国 LoRaWAN 网络用于 TAIS 及其他服务。"],
    ["现有网关", "客户已部署 Milesight Solar Gateway SG50-L08GL-868M PN 30W。"],
    ["基础硬件路线", "采用 S275 为基础设备，将原 4G 模块替换为 LoRaWAN EU868 模块，集成大容量锂电池至防护外壳。"],
    ["集成范围", "S275 主板、LoRaWAN 模块、锂电池、太阳能供电/充电、继电器输出、传感器端子、外置天线接口、防护外壳及客户 Logo。"],
    ["本地报警要求", "终端必须具备本地逻辑（IF-THEN）。无网络连接时，传感器触发后仍需立即驱动警笛/警灯，不依赖云平台指令。"],
    ["样机数量", "初始测试样机 20 台。我方先提供 2 台工程样机。"],
])

# ====== SECTION 2 ======
sec_title(doc, "二", "关键硬件与 I/O 接口需求")
mk_table(doc, ["模块", "规格说明"], [
    ["LoRaWAN 通信模块", "EU868（863–870 MHz）；8 个 multi-SF 通道；支持 ADR。"],
    ["数据链路", "节点 → Milesight SG50 网关 → 应用服务器 → MQTT 云平台。"],
    ["数字输入 DI", "8 路干接点输入：门磁、PIR、振动、电源状态、防拆回路。"],
    ["模拟量输入 AI", "2 路，支持 4–20 mA 或 0–10 V，12 位精度，用于油位/压力监测。"],
    ["温度监测", "变压器油温/绕组温度，通过 PT100 转 4–20 mA 变送器接入。"],
    ["继电器输出 DO", "2 路：DO1=警笛/警灯，DO2=远程负荷切断/联锁。额定 5A @ 250VAC / 30VDC。"],
    ["本地逻辑引擎", "IF-THEN 逻辑：PIR、门磁、振动、断电、温度、阀门破坏均可触发。"],
    ["外置天线", "8 dBi 玻璃钢天线，含 3 米馈线。"],
    ["配置工具", "PC 端配置软件，功能等同或优于 Milesight UC300。"],
    ["品牌标识", "最终产品外壳需印制客户 Logo。"],
])

# ====== SECTION 3 ======
sec_title(doc, "三", "传感器与输出功能清单")
mk_table(doc, ["功能", "说明"], [
    ["门/检修口监测", "磁簧开关，用于柜门/检修口开启监测。"],
    ["人体移动检测（PIR）", "检测变压器站点附近人员活动。"],
    ["振动检测", "加速度计触发，检测物理破坏或变压器移动。"],
    ["市电状态监测", "监测交流供电，检测断电（盗窃指示信号）。"],
    ["紧急/防拆回路", "阀门破坏检测，用于防盗油。"],
    ["PT100 温度", "油温和绕组温度监测（通过 PT100 转 4–20mA 变送器）。"],
    ["油位/油压", "通过 AI 接口连接 4–20 mA 传感器监测。"],
    ["警笛/警灯（DO1）", "本地报警 — 检测到入侵时即时触发。"],
    ["负荷切断（DO2）", "远程负荷切除或系统联锁控制。"],
    ["RS485 继电器板", "兼容客户现有 8 通道 RS485 继电器扩展板。"],
    ["物联网锁", "电子柜锁（含隐藏机械应急钥匙）。集成 ThingsBoard 平台，支持远程开锁、操作人员记录及维护审计追踪。"],
])

# ====== SECTION 4 ======
sec_title(doc, "四", "供电、电池、太阳能与机壳")
mk_table(doc, ["项目", "规格说明"], [
    ["电池续航", "断电期间至少运行 72 小时（含所有传感器及警笛工作）。"],
    ["电池类型", "11.1V（标称 3S 锂离子）/ 12V 兼容。5Ah 可充电，带 BMS 保护。"],
    ["主供电方式", "太阳能一体化方案：太阳能板 + 锂电池。避免变压器取电的浪涌损害。"],
    ["太阳能输入", "DC 9–36 V（S275 输入范围）。我方提供太阳能板规格参数。"],
    ["安装方式", "（1）户外杆装：太阳能板固定在 TAIS 单元上，配杆装支架；\n（2）室内壁挂：太阳能板安装于屋顶，通过线缆连接至设备。"],
    ["充电管理", "100 W MPPT 智能充电器。含 BMS、状态检测、低压告警、过充保护。"],
    ["机壳", "专业防护外壳，支持 DIN 导轨/壁挂安装，紧凑设计。"],
    ["户外防护", "防水、防尘、防腐蚀，适用于所有外部连接。"],
])

# ====== SECTION 5 ======
sec_title(doc, "五", "平台系统需求（ThingsBoard 边界）")
body_txt(doc, "说明：平台部分独立于 S275 硬件。此处仅涉及 payload 格式、数据字段、告警规则及集成接口参考，供技术团队提前知悉。", sz=6.5, clr="B91C1C", it=True)
mk_table(doc, ["项目", "需求说明"], [
    ["通信链路", "MQTT 云平台，与 Milesight 网关、LoRaWAN 终端节点、应用服务器通信。"],
    ["告警通知", "WhatsApp、短信、电子邮件。未来可扩展至电话通知警方或快速响应团队。"],
    ["地图仪表盘", "控制中心可在地图上查看所有站点：绿色=在线，红色=离线；报警时显示具体报警站点。"],
    ["前后端开发", "完整前端+后端开发，替代当前昂贵且效果不佳的 Akenza API 方案。"],
    ["物联网锁集成", "基于 ThingsBoard 平台：远程开锁、授权操作人员记录、维护审计追踪。"],
    ["集成边界", "平台范围、接口、开发周期及费用独立定义。硬件项目需提供清晰的 payload 定义、设备数据字段及联调支持。"],
])

# ====== SECTION 6 ======
sec_title(doc, "六", "集成配件与商务方案")
dark_banner(doc,
    "开发与工程 NRE 费用：10,000.00 美元（All-Inclusive）",
    "涵盖：定制固件开发、LoRaWAN 模块集成、太阳能充电管理设计、物联网锁接线、2 台完整功能样机及空运至哈拉雷运费。")
doc.add_paragraph().paragraph_format.space_after = Pt(2)

mk_table(doc, ["数量", "太阳能板", "定制主机 (USD)", "外购配件* (USD)", "总价 (EXW)"], [
    ["2,000", "30 W", "$80.00", "$36.63", "$116.63"],
    ["2,000", "40 W", "$80.00", "$39.13", "$119.13"],
    ["2,000", "50 W", "$80.00", "$32.83", "$112.83"],
    ["5,000", "30 W", "$70.00", "$36.63", "$106.63"],
    ["5,000", "40 W", "$70.00", "$39.13", "$109.13"],
    ["5,000", "50 W", "$70.00", "$32.83", "$102.83"],
    ["10,000", "30 W", "$65.00", "$35.85", "$100.85"],
    ["10,000", "40 W", "$65.00", "$38.35", "$103.35"],
    ["10,000", "50 W", "$65.00", "$32.05", "$97.05"],
])

body_txt(doc, "*外购配件清单：11.1V / 5Ah 锂电池、100W MPPT 太阳能充电器、物联网柜锁（隐藏应急钥匙）、8dBi 天线+3米馈线、杆装支架。", sz=6, clr="B91C1C", it=True, sb=2, sa=1)
body_txt(doc, "50W 促销优惠价：2,000 台 @ $149.00  |  5,000 台 @ $144.00  |  10,000 台 @ $139.00", sz=7.5, clr="1F4E79", b=True)

# Payment Terms
body_txt(doc, "", sz=2, sa=2)
p_pt = doc.add_paragraph()
p_pt.paragraph_format.space_before = Pt(1); p_pt.paragraph_format.space_after = Pt(1)
r_pt_title = p_pt.add_run("付款条款")
r_pt_title.bold = True; r_pt_title.font.size = Pt(10); r_pt_title.font.color.rgb = RGBColor.from_string("0F172A"); rf(r_pt_title, ea="微软雅黑")

pt_tbl = doc.add_table(rows=3, cols=2); pt_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(pt_tbl)
for j, h in enumerate(["条款", "条件"]):
    c = pt_tbl.rows[0].cells[j]; mg(c, 50, 50, 100, 100)
    set_cell_border(c, "bottom", 10, "1F4E79")
    cll(c, h, b=True, sz=8.5, clr="1F4E79")
pt_data = [
    ("定金", "订单确认后支付总金额的 50%。"),
    ("尾款", "发货前付清剩余 50%。"),
]
for i, (l, v) in enumerate(pt_data):
    cl = pt_tbl.rows[i+1].cells[0]; mg(cl, 40, 40, 100, 100)
    cll(cl, l, b=True, sz=8.5, clr="1E293B")
    cv = pt_tbl.rows[i+1].cells[1]; mg(cv, 40, 40, 100, 100)
    cll(cv, v, sz=8.5, clr="475569")

body_txt(doc, "报价术语：EXW（深圳）。所有价格以美元计。", sz=6.5, clr="B91C1C", it=True, sb=1, sa=0)

# ====== SECTION 7 ======
sec_title(doc, "七", "项目开发与交付时间表")
body_txt(doc, "初步排期 — 以规格冻结及定金支付后确认为准。", sz=6.5, clr="B91C1C", it=True)
mk_table(doc, ["阶段", "时间/周期"], [
    ["原理图与 PCB 硬件修改", "规格冻结并支付定金后 4 周"],
    ["样机生产与调试（2 台）", "15 个工作日（含固件移植）"],
    ["样机交付（哈拉雷）", "国际快递空运 — 待定"],
    ["首批量产（2,000 台）", "样机验收通过后 30 个工作日"],
    ["月产能", "400–700 台/月（持续稳定，可扩展）"],
])

# ====== SECTION 8 ======
sec_title(doc, "八", "客户确认与签字")
info_box(doc,
    '“本人/我公司确认已审阅本文件。文件中所列需求、技术规格及项目理解在签署之日均为完整且准确的。”')
doc.add_paragraph().paragraph_format.space_after = Pt(4)

st = doc.add_table(rows=7, cols=2); st.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(st)
for j, h in enumerate(["填写项目", "客户填写内容"]):
    c = st.rows[0].cells[j]; mg(c, 60, 60, 100, 100)
    set_cell_border(c, "bottom", 10, "1F4E79")
    cll(c, h, b=True, sz=8.5, clr="1F4E79")
sig_data = ["客户公司名称", "授权代表姓名", "职务/职称", "签字", "日期"]
for i, lbl in enumerate(sig_data):
    rn = i+1
    cl = st.rows[rn].cells[0]; cv = st.rows[rn].cells[1]
    mg(cl, 50, 50, 100, 100); cll(cl, lbl, b=True, sz=8.5, clr="1E293B")
    mg(cv, 50, 50, 100, 100); cll(cv, "_"*50, sz=7, clr="CBD5E1")

sr = st.rows[6]; sr.cells[0].merge(sr.cells[1])
mg(sr.cells[0], 50, 50, 100, 100)
set_cell_border(sr.cells[0], "top", 4, "CBD5E1")
cll(sr.cells[0], "公司公章（如适用）", sz=8, clr="64748B")

# --- END ---
pe = doc.add_paragraph(); pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
pe.paragraph_format.space_before = Pt(16)
re = pe.add_run("— 文档结束 —")
re.font.size = Pt(9); re.font.color.rgb = RGBColor.from_string("94A3B8"); re.italic = True; rf(re)

doc.save(OUT)
print("OK: " + OUT)
