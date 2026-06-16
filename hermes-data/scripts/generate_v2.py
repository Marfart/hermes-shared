#!/c/Users/Admin/AppData/Local/Programs/Python/Python313/python.exe
# -*- coding: utf-8 -*-
"""Generate professional document with proper watermark centered on page."""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree as ET
import copy

LOGO_PATH = r"C:\Users\Admin\Desktop\Working\津巴布尔客户\签署文件\钡铼英文LOGO.png"
OUTPUT_PATH = r"C:\Users\Admin\AppData\Local\hermes\更新后新文件.docx"

doc = Document()

# ===== Page Setup =====
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ========================================================
# PROPER WATERMARK - simplified approach
# ========================================================
def add_proper_watermark(section, image_path):
    """Add logo as centered watermark behind text on every page"""
    header = section.header
    header.is_linked_to_previous = False
    
    # Clear existing content
    for p in header.paragraphs:
        p.clear()
    
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    # Center alignment
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run()
    # Add at a larger size so it fills the middle area
    run.add_picture(image_path, width=Inches(3.5), height=Inches(2.2))
    
    # Now convert inline to anchored behind-text
    drawing = run._r.find(qn('w:drawing'))
    if drawing is None:
        return
    
    inline = drawing.find(qn('wp:inline'))
    if inline is None:
        return
    
    # Get extent and graphic content
    extent = inline.find(qn('wp:extent'))
    cx = extent.get('cx') if extent is not None else '3200400'
    cy = extent.get('cy') if extent is not None else '2009000'
    
    graphic = inline.find(qn('a:graphic'))
    
    # Calculate center position: half page minus half image
    # Letter page: 8.5in x 11in = 7772400 x 10058400 EMUs
    # A4 page: 21cm x 29.7cm = about 7938000 x 11226600 EMUs
    # Page width in EMU approximately 7938000 (A4 with 2.54cm margins = width ~16cm = ~6048000 EMU)
    # Actually, let's use simpler positioning
    
    # Build anchor XML with all required namespace declarations
    anchor_xml = '''<wp:anchor
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        distT="0" distB="0" distL="0" distR="0"
        simplePos="0" relativeHeight="0"
        behindDoc="1" locked="1"
        layoutInCell="0" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page">
            <wp:posOffset>2100000</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="page">
            <wp:posOffset>4200000</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="''' + str(cx) + '''" cy="''' + str(cy) + '''"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="999" name="Watermark" descr="Company Logo Background"/>
        <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
    </wp:anchor>'''
    
    anchor = ET.fromstring(anchor_xml.encode('utf-8'))
    
    # Copy the graphic element into the anchor
    if graphic is not None:
        anchor.append(copy.deepcopy(graphic))
    
    # Replace inline with anchor
    drawing.remove(inline)
    drawing.append(anchor)


add_proper_watermark(doc.sections[0], LOGO_PATH)

# ========================================================
# HELPER FUNCTIONS
# ========================================================

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=80, bottom=80, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(existing)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_run_font(run, name="Calibri", east_asia="微软雅黑"):
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{east_asia}"/>')
    rPr.append(rFonts)

def add_cell_para(cell, text, bold=False, font_size=10, color="1E293B", 
                   alignment=None, line_spacing=1.15, space_after=0):
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = line_spacing
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    set_run_font(run)
    return p

def add_section_title(doc, number, title):
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_after = Pt(2)
    p_sep.paragraph_format.space_before = Pt(14)
    p_sep.paragraph_format.line_spacing = 0.3
    run_sep = p_sep.add_run("━" * 60)
    run_sep.font.size = Pt(1)
    run_sep.font.color.rgb = RGBColor.from_string("93C5FD")
    set_run_font(run_sep)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    
    run_num = p.add_run(f"{number}. ")
    run_num.bold = True
    run_num.font.size = Pt(14)
    run_num.font.color.rgb = RGBColor.from_string("1D4ED8")
    set_run_font(run_num)
    
    run_title = p.add_run(f"{title}")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor.from_string("0F172A")
    set_run_font(run_title)

def add_table_borders(table, top_color="CBD5E1", bottom_color="CBD5E1", 
                      left_color="E2E8F0", right_color="E2E8F0",
                      inside_h_color="E2E8F0", inside_v_color="E2E8F0"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{top_color}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{bottom_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{left_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{right_color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{inside_h_color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{inside_v_color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def remove_all_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    
    # Header row - dark navy
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1E3A5F")
        set_cell_margins(cell, 100, 100, 150, 150)
        add_cell_para(cell, header, bold=True, font_size=9, color="FFFFFF")
    
    # Data rows
    for r, row_data in enumerate(rows):
        for c, text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            bg = "F8FAFC" if r % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)
            set_cell_margins(cell, 70, 70, 150, 150)
            is_label = (c == 0)
            add_cell_para(cell, str(text), bold=is_label, font_size=8.5, 
                         color="1E293B" if is_label else "475569")
    return table

def add_body_text(doc, text, font_size=9, color="334155", bold=False, italic=False, space_before=0, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    set_run_font(run)
    return p

def add_info_box(doc, text, font_size=9, text_color="1E3A5F"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EFF6FF")
    set_cell_margins(cell, 120, 120, 180, 180)
    add_table_borders(table, top_color="BFDBFE", bottom_color="BFDBFE", 
                      left_color="BFDBFE", right_color="BFDBFE",
                      inside_h_color="BFDBFE", inside_v_color="BFDBFE")
    add_cell_para(cell, text, font_size=font_size, color=text_color, line_spacing=1.2)


# ========================================================
# BUILD DOCUMENT
# ========================================================

# ===== HEADER =====
title_table = doc.add_table(rows=1, cols=2)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_all_borders(title_table)

# Add bottom border
tbl = title_table._tbl
tblPr = tbl.tblPr
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="24" w:space="0" w:color="1E3A5F"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

# Title
title_cell = title_table.rows[0].cells[0]
set_cell_margins(title_cell, 0, 0, 0, 200)
p1 = title_cell.paragraphs[0]
p1.paragraph_format.space_after = Pt(2)
run1 = p1.add_run("Zimbabwe ZODSAT\nTransformer Anti-Theft & Intrusion Prevention")
run1.bold = True
run1.font.size = Pt(20)
run1.font.color.rgb = RGBColor.from_string("0F172A")
set_run_font(run1)

p2 = title_cell.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
p2.paragraph_format.space_before = Pt(2)
run2 = p2.add_run("S275 LoRaWAN Custom Terminal Node — Product Specification & Requirements Confirmation Document")
run2.bold = True
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor.from_string("1D4ED8")
set_run_font(run2)

# Logo
logo_cell = title_table.rows[0].cells[1]
set_cell_margins(logo_cell, 0, 0, 0, 0)
pl = logo_cell.paragraphs[0]
pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_logo = pl.add_run()
run_logo.add_picture(LOGO_PATH, width=Inches(1.5), height=Inches(0.94))

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ===== INFO TABLE =====
info_data = [
    ("Prepared for", "ZODSAT (Zimbabwe)"),
    ("Prepared by", "Shenzhen Beilai Technology Co., Ltd. (BLIIOT)"),
    ("Document Purpose", "Detailed product specification for customer review — all features, functions, design scope, accessories, and project delivery timeline."),
    ("Document Version", "V2.0 — Updated Specification"),
    ("Date", "27 May 2026"),
    ("Classification", "Confidential — For ZODSAT Review Only"),
]

info_table = doc.add_table(rows=len(info_data), cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(info_table)

for i, (label, value) in enumerate(info_data):
    bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
    cl = info_table.rows[i].cells[0]
    cv = info_table.rows[i].cells[1]
    set_cell_shading(cl, bg)
    set_cell_margins(cl, 60, 60, 150, 150)
    add_cell_para(cl, label, bold=True, font_size=9, color="1E293B")
    set_cell_shading(cv, "FFFFFF" if i % 2 == 0 else "F8FAFC")
    set_cell_margins(cv, 60, 60, 150, 150)
    add_cell_para(cv, value, bold=False, font_size=9, color="475569")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ===== CUSTOMER REVIEW INSTRUCTIONS =====
add_info_box(doc, 
    "📋 Customer Review Instructions\n\n"
    "This document presents the complete technical specifications, product features, design scope, "
    "integrated accessories, and project timeline for the Zimbabwe ZODSAT Transformer Anti-Theft "
    "and Intrusion Prevention Project, based on our technical evaluations and discussions.\n\n"
    "Please review every item carefully. If any information is missing, inaccurate, or inconsistent "
    "with your expectations, please mark it in this document or notify us in writing.\n\n"
    "If all information is correct, please sign the Customer Confirmation section at the end. "
    "The signed version will serve as the basis for prototype design, quotation, and project execution."
)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ========================================================
# SECTION 1
# ========================================================
add_section_title(doc, "1", "Overall Technical Route")

make_table(doc, 
    ["Item", "Requirement / Understanding"],
    [
        ["Device Type", "LoRaWAN End Node / Terminal Node. BLIIOT does not need to provide a LoRaWAN Gateway."],
        ["Communication", "LoRaWAN only (no 4G). Customer's 4G network is unstable; previous system had alarm delays. Customer is building a national LoRaWAN network for TAIS."],
        ["Existing Gateway", "Milesight Solar Gateway SG50-L08GL-868M PN 30W — already deployed."],
        ["Base Hardware", "BLIIOT S275 base device with LoRaWAN EU868 module replacing the original 4G module. Large-capacity lithium battery in professional protective enclosure."],
        ["Integrated Scope", "S275 mainboard, LoRaWAN module, lithium battery, solar charging management, relay outputs, sensor terminals, external antenna interface, protective enclosure, and customer logo."],
        ["Local Alarm Logic", "Terminal MUST support IF-THEN local logic. Siren/strobe activates immediately on sensor trigger even without LoRaWAN — NO cloud dependency."],
        ["Prototype Qty", "20 test prototypes. BLIIOT to provide 2 engineering prototypes first for evaluation."],
    ]
)

# ========================================================
# SECTION 2
# ========================================================
add_section_title(doc, "2", "Key Hardware & I/O Requirements")

make_table(doc,
    ["Module", "Specification"],
    [
        ["LoRaWAN Module", "EU868 (863–870 MHz); 8 multi-SF channels (868.1–867.9 MHz); ADR supported."],
        ["Data Path", "Node → Milesight SG50 Gateway → Application Server → MQTT Cloud Platform."],
        ["Digital Input (DI)", "8 Dry Contact inputs: door/hatch magnetic switches, PIR, vibration, power status, tamper loop."],
        ["Analog Input (AI)", "2 channels, 4–20 mA or 0–10 V, 12-bit resolution. For oil level / pressure."],
        ["Temperature", "Transformer oil/winding temp via PT100-to-4–20mA transmitter solution."],
        ["Relay Output (DO)", "2 relays: DO1 = siren/strobe, DO2 = load shedding/interlocking. 5A @ 250 VAC / 30 VDC."],
        ["Local Logic Engine", "IF-THEN engine. Triggers: PIR, door switch, vibration, power loss, temperature, valve tamper."],
        ["External Antenna", "7 dBi antenna with 3-metre cable for flexible placement outside transformer rooms."],
        ["Configuration Tool", "PC-based software — equivalent or superior to Milesight UC300."],
        ["Branding", "Customer logo on final product enclosure."],
    ]
)

# ========================================================
# SECTION 3
# ========================================================
add_section_title(doc, "3", "Sensor & Output Function List")

make_table(doc,
    ["Function", "Description"],
    [
        ["Door / Access Hatch Sensor", "Magnetic switches for enclosure door / hatch monitoring."],
        ["PIR Motion Detection", "Human presence detection near transformer sites."],
        ["Vibration Sensor", "Accelerometer-based trigger for physical tampering or transformer movement."],
        ["Mains Power Status", "AC mains monitoring to detect power loss (theft indicator)."],
        ["Emergency / Tamper Loop", "Valve tampering detection for oil theft prevention."],
        ["PT100 Temperature", "Oil and winding temperature via PT100-to-4–20 mA transmitter."],
        ["Oil Level / Pressure", "AI-based monitoring via 4–20 mA sensors."],
        ["Siren / Strobe (DO1)", "Local alarm — triggers instantly on breach detection."],
        ["Load Shedding (DO2)", "Remote load shedding or system interlocking."],
        ["RS485 Relay Board", "Compatible with existing 8-channel RS485 relay board."],
    ]
)

# ========================================================
# SECTION 4
# ========================================================
add_section_title(doc, "4", "Power Supply, Battery, Solar & Enclosure")

make_table(doc,
    ["Item", "Specification"],
    [
        ["Battery Backup", "Min. 72 hours during power outage (all sensors + siren active). 5 Ah Li-ion provides ~4 days without solar."],
        ["Battery Type", "12 V / 5 Ah rechargeable lithium battery pack with BMS (overcharge/discharge protection)."],
        ["Primary Power", "Turnkey solar: solar panel + lithium battery. Replaces transformer-tap power (eliminates surge damage and SPD/OVP costs)."],
        ["Solar Input Range", "DC 9–36 V (S275 input range). BLIIOT to provide panel specs."],
        ["Installation", "Option 1: Outdoor / pole — panel on TAIS unit. Option 2: Indoor / wall — panel on roof, cabled to unit."],
        ["Charging", "100 W MPPT intelligent solar charger. BMS, status detection, low-voltage alarm, overcharge/discharge protection."],
        ["Enclosure", "Professional protective enclosure, DIN-rail or wall-mount. Compact size for mainboard, battery, relays."],
        ["Outdoor Protection", "Waterproof, dustproof, corrosion-resistant for all external connections (antenna, sensors, siren, solar cable)."],
    ]
)

# ========================================================
# SECTION 5
# ========================================================
add_section_title(doc, "5", "Platform System Requirements (ThingsBoard Boundary)")

add_body_text(doc, 
    "Note: Platform requirements are separate from the S275 hardware scope but affect payload format, "
    "data fields, alarm rules, and integration interfaces. Included here for reference.",
    font_size=8.5, color="64748B", italic=True, space_after=6)

make_table(doc,
    ["Item", "Requirement"],
    [
        ["Communication", "MQTT cloud platform integrating with Milesight Gateway, LoRaWAN End Nodes, and Application Server."],
        ["Alarm Notifications", "WhatsApp, SMS, Email. Future: phone call to police / rapid-response teams."],
        ["Map Dashboard", "Site map: green = online, red = offline. Alarm events show site location."],
        ["Front-end & Back-end", "Full UI + backend development to replace Akenza API."],
        ["Boundary", "Platform scope, timeline, cost, and responsibility defined separately. Hardware provides payload definitions, data fields, and integration support."],
    ]
)

# ========================================================
# SECTION 6
# ========================================================
add_section_title(doc, "6", "Integrated Accessories & Commercial Scheme")

# NRE box
nre_table = doc.add_table(rows=1, cols=1)
nre_table.alignment = WD_TABLE_ALIGNMENT.CENTER
nre_cell = nre_table.rows[0].cells[0]
set_cell_shading(nre_cell, "1E3A5F")
set_cell_margins(nre_cell, 100, 100, 180, 180)
add_table_borders(nre_table, top_color="1E3A5F", bottom_color="1E3A5F", 
                  left_color="1E3A5F", right_color="1E3A5F",
                  inside_h_color="1E3A5F", inside_v_color="1E3A5F")
add_cell_para(nre_cell, "Development & Engineering NRE Fee: USD $10,000.00 (All-Inclusive)", 
             bold=True, font_size=11, color="FFFFFF", alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_para(nre_cell, 
    "Covers: custom firmware, LoRaWAN module integration, solar charging design, "
    "IoT lock wiring, 2 prototype samples, and express air freight to Harare, Zimbabwe.",
    font_size=8.5, color="93C5FD", alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Pricing table
make_table(doc,
    ["Qty", "Solar Panel", "Base Node (USD)", "Accessories* (USD)", "Total (EXW)"],
    [
        ["2,000", "30 W", "$80.00", "$36.63", "$116.63"],
        ["2,000", "40 W", "$80.00", "$39.13", "$119.13"],
        ["2,000", "50 W", "$80.00", "$43.13", "$123.13"],
        ["5,000", "30 W", "$70.00", "$36.63", "$106.63"],
        ["5,000", "40 W", "$70.00", "$39.13", "$109.13"],
        ["5,000", "50 W", "$70.00", "$43.13", "$113.13"],
        ["10,000", "30 W", "$65.00", "$35.85", "$100.85"],
        ["10,000", "40 W", "$65.00", "$38.35", "$103.35"],
        ["10,000", "50 W", "$65.00", "$42.35", "$107.35"],
    ]
)

add_body_text(doc, 
    "*Accessories include: 12 V / 5 Ah Li-ion Battery Pack, 100 W MPPT Solar Charger, "
    "12 V DC IoT Cabinet Lock (hidden backup key), 8 dBi Fibreglass Antenna + 3 m Cable, "
    "and Heavy-duty Pole-Mount Stand.",
    font_size=8, color="64748B", italic=True, space_before=2, space_after=1)

add_body_text(doc,
    "Promotional pricing (50 W): 2,000 pcs @ $149.00 | 5,000 pcs @ $144.00 | 10,000 pcs @ $139.00",
    font_size=8.5, color="1D4ED8", bold=True, space_before=0, space_after=2)

# ========================================================
# SECTION 7
# ========================================================
add_section_title(doc, "7", "Project Development & Delivery Schedule")

add_body_text(doc, 
    "Note: The timeline below is preliminary and subject to final confirmation upon specification "
    "freeze and deposit. Production durations will be finalised once accessory selections are confirmed.",
    font_size=8.5, color="B91C1C", italic=True, space_after=6)

make_table(doc,
    ["Phase", "Timeline"],
    [
        ["Schematic & PCB Hardware Modification", "4 weeks from specification freeze and deposit"],
        ["Prototype Production & Debugging (2 units)", "15 working days (incl. firmware porting)"],
        ["Prototype Delivery (Harare)", "International express air freight — TBC with logistics partner"],
        ["Mass Production (first 2,000 pcs)", "30 working days after prototype approval"],
        ["Monthly Production Capacity", "400–700 units / month (sustained, scalable)"],
    ]
)

# ========================================================
# SECTION 8
# ========================================================
add_section_title(doc, "8", "Customer Confirmation & Signature")

add_info_box(doc,
    '"I / We confirm that we have reviewed this document. The requirements, technical specifications, '
    'and project understanding are complete and accurate as of the signing date. This signed confirmation '
    'may be used as the basis for prototype design, quotation, and project execution."',
    font_size=9, text_color="1E293B"
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Signature table
sig_table = doc.add_table(rows=7, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_borders(sig_table)

sig_fields_data = [
    ("Field", "Information to be Completed by Customer"),
    ("Customer Company", None),
    ("Authorized Representative", None),
    ("Position / Title", None),
    ("Signature", None),
    ("Date", None),
]

for i, (label, val) in enumerate(sig_fields_data):
    if i == 0:
        for j, h in enumerate([label, val or label]):
            cell = sig_table.rows[0].cells[j]
            set_cell_shading(cell, "1E3A5F")
            set_cell_margins(cell, 100, 100, 150, 150)
            add_cell_para(cell, h, bold=True, font_size=9, color="FFFFFF")
    else:
        cl = sig_table.rows[i].cells[0]
        cv = sig_table.rows[i].cells[1]
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        set_cell_shading(cl, bg)
        set_cell_margins(cl, 80, 80, 150, 150)
        add_cell_para(cl, label, bold=True, font_size=9, color="1E293B")
        set_cell_margins(cv, 80, 80, 150, 150)
        add_cell_para(cv, "_" * 45, font_size=9, color="CBD5E1")

# Stamp row
stamp_cell = sig_table.rows[6].cells[0]
stamp_cell.merge(sig_table.rows[6].cells[1])
set_cell_shading(stamp_cell, "FAFAFA")
set_cell_margins(stamp_cell, 100, 100, 150, 150)
add_cell_para(stamp_cell, "Company Stamp / Seal (if applicable)", font_size=9, color="64748B")

# ========================================================
# END MARKER
# ========================================================
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(16)
run_end = p_end.add_run("—— End of Document ——")
run_end.font.size = Pt(10)
run_end.font.color.rgb = RGBColor.from_string("94A3B8")
run_end.font.italic = True
set_run_font(run_end)

# ========================================================
# SAVE
# ========================================================
doc.save(OUTPUT_PATH)
print(f"✅ Document saved to: {OUTPUT_PATH}")
